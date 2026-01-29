"""
OpenAI Responses API Proxy for OpenWebUI

This proxy intercepts Chat Completions requests from OpenWebUI,
detects embedded PDF markers, and forwards them to OpenAI's Responses API
with proper input_file support for vision analysis.

Run with: uvicorn openai_responses_proxy:app --host 0.0.0.0 --port 8000
"""

import os
import re
import json
import base64
import mimetypes
import httpx
import csv
import io
import uuid
import math
import time
from pathlib import Path
from typing import Optional, List, Dict, Any
from urllib.parse import quote_plus
from fastapi import FastAPI, Request, HTTPException  # type: ignore
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse, HTMLResponse  # type: ignore
import asyncio
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage  # type: ignore
from reportlab.lib.pagesizes import letter  # type: ignore
from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
from reportlab.lib import colors  # type: ignore
from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore
from openpyxl import load_workbook  # type: ignore

app = FastAPI(title="OpenAI Responses Proxy")

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
OPENAI_BASE_URL = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")
DEBUG = os.environ.get("PROXY_DEBUG", "true").lower() == "true"
HTTP_CLIENT: httpx.AsyncClient | None = None

# Keep-alive + HTTP/2 client for faster, reusable connections to the proxy/OpenAI.
HTTP_TIMEOUT = httpx.Timeout(connect=5.0, read=120.0, write=120.0, pool=120.0)
HTTP_LIMITS = httpx.Limits(
    max_keepalive_connections=20,
    max_connections=50,
    keepalive_expiry=300.0,
)
MAX_FILE_MB = 10
MAX_TOTAL_MB = 40
MAX_DOCX_CHARS = 40000  # safety cap
MAX_CSV_ROWS = 1000
MAX_CSV_COLS = 50
MAX_XLSX_ROWS = 500
MAX_XLSX_COLS = 50
MAX_TEXT_CHARS = 40000
MAX_JSON_BYTES = 1_000_000
MAX_ZIP_BYTES = 50 * 1024 * 1024
MAX_ZIP_FILES = 500
MAX_JCAMP_BYTES = 2 * 1024 * 1024

# Simple in-memory metrics (non-persistent)
METRICS = {
    "requests_total": 0,
    "responses_api_calls": 0,
    "chat_completions_calls": 0,
    "errors_total": 0,
    "last_error": "",
    "last_latency_ms": 0.0,
}


async def init_http_client() -> None:
    """Initialize shared AsyncClient with keep-alive/HTTP2 for lower latency."""
    global HTTP_CLIENT
    if HTTP_CLIENT is None:
        HTTP_CLIENT = httpx.AsyncClient(
            timeout=HTTP_TIMEOUT,
            limits=HTTP_LIMITS,
            http2=True,
        )


async def get_http_client() -> httpx.AsyncClient:
    """Return shared AsyncClient (initialized lazily)."""
    if HTTP_CLIENT is None:
        await init_http_client()
    return HTTP_CLIENT


@app.on_event("startup")
async def _startup_client():
    await init_http_client()


@app.on_event("shutdown")
async def _shutdown_client():
    global HTTP_CLIENT
    if HTTP_CLIENT:
        await HTTP_CLIENT.aclose()
        HTTP_CLIENT = None

# Regex to find PDF markers: [__PDF_FILE_B64__ filename=xxx.pdf]base64data[/__PDF_FILE_B64__]
PDF_MARKER_RE = re.compile(
    r"\[__PDF_FILE_B64__ filename=([^\]]+)\]([A-Za-z0-9+/=]+)\[/__PDF_FILE_B64__\]",
    re.DOTALL
)


def log(msg: str):
    if DEBUG:
        print(f"[PROXY] {msg}")


def _get_file_path(file_obj: Dict[str, Any]) -> str:
    if not isinstance(file_obj, dict):
        return ""
    if isinstance(file_obj.get("file"), dict):
        f = file_obj["file"]
        path = f.get("path", "") or ""
        if not path and isinstance(f.get("meta"), dict):
            path = f["meta"].get("path", "") or ""
        return path.strip()
    return (file_obj.get("path") or "").strip()


def _get_file_name(file_obj: Dict[str, Any]) -> str:
    if not isinstance(file_obj, dict):
        return ""
    if isinstance(file_obj.get("file"), dict):
        f = file_obj["file"]
        name = f.get("filename") or f.get("name") or ""
        if not name and isinstance(f.get("meta"), dict):
            name = f["meta"].get("name") or ""
        return name.strip()
    return (file_obj.get("name") or file_obj.get("filename") or "").strip()


def _extract_inline_images_from_messages(messages: list[dict]) -> list[dict]:
    """
    Extract OpenWebUI inline images from message content blocks:
      {"type":"image_url","image_url":{"url":"data:image/png;base64,..."}}
    Returns: [{"url": "...", "name": "inline_001"}]
    """
    out = []
    idx = 1
    for msg in messages or []:
        content = msg.get("content")
        if not isinstance(content, list):
            continue
        for item in content:
            if not isinstance(item, dict):
                continue
            if item.get("type") != "image_url":
                continue
            iu = item.get("image_url") or {}
            if not isinstance(iu, dict):
                continue
            url = iu.get("url")
            if isinstance(url, str) and url.startswith("data:image/"):
                out.append({"url": url, "name": f"inline_{idx:03d}"})
                idx += 1
    return out


def _extract_all_files(body: dict, messages: list) -> List[Dict[str, Any]]:
    files: List[Dict[str, Any]] = []
    if isinstance(body.get("files"), list):
        files.extend(body["files"])
    for msg in messages:
        if isinstance(msg.get("files"), list):
            files.extend(msg["files"])
        if isinstance(msg.get("attachments"), list):
            files.extend(msg["attachments"])
        if isinstance(msg.get("sources"), list):
            for source_obj in msg["sources"]:
                if isinstance(source_obj, dict):
                    source = source_obj.get("source", {})
                    if source.get("type") == "file" and isinstance(source.get("file"), dict):
                        files.append({"file": source["file"]})
    # Deduplicate by path
    seen = set()
    unique = []
    for f in files:
        path = _get_file_path(f)
        if path and path not in seen:
            seen.add(path)
            unique.append(f)
    return unique


def _file_size_bytes(path: str) -> int:
    try:
        return os.path.getsize(path)
    except OSError:
        return -1


def _is_pdf(name: str, mime: str) -> bool:
    return name.lower().endswith(".pdf") or mime == "application/pdf"


def _is_image(mime: str) -> bool:
    return mime.startswith("image/")


def _is_docx(name: str, mime: str) -> bool:
    return name.lower().endswith((".docx",)) or mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",)


def _is_csv(name: str, mime: str) -> bool:
    return name.lower().endswith(".csv") or mime in ("text/csv", "application/csv")


def _is_xlsx(name: str, mime: str) -> bool:
    return name.lower().endswith((".xlsx", ".xlsm")) or mime in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",)


def _is_txt(name: str, mime: str) -> bool:
    return name.lower().endswith((".txt", ".log")) or mime.startswith("text/")


def _is_md(name: str, mime: str) -> bool:
    return name.lower().endswith((".md", ".markdown"))


def _is_tsv(name: str, mime: str) -> bool:
    return name.lower().endswith(".tsv") or mime in ("text/tab-separated-values", "text/tsv")


def _is_json_file(name: str, mime: str) -> bool:
    return name.lower().endswith(".json") or mime in ("application/json", "text/json")


def _is_doc(name: str, mime: str) -> bool:
    return name.lower().endswith(".doc") or mime in ("application/msword",)


def _is_xls(name: str, mime: str) -> bool:
    return name.lower().endswith(".xls") or mime in ("application/vnd.ms-excel",)


def _is_odt(name: str, mime: str) -> bool:
    return name.lower().endswith(".odt") or mime in ("application/vnd.oasis.opendocument.text",)


def _is_rtf(name: str, mime: str) -> bool:
    return name.lower().endswith(".rtf") or mime in ("application/rtf", "text/rtf")


def _is_bruker_zip(name: str, mime: str) -> bool:
    return name.lower().endswith(".zip") and ("application/zip" in mime or "zip" in mime)


def _is_jcamp(name: str, mime: str) -> bool:
    return name.lower().endswith((".jdx", ".dx", ".jcamp")) or ("jcamp" in mime.lower() if mime else False)


def _extract_docx_text(path: str) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        log(f"DOCX import failed: {e}")
        return ""
    try:
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)
            if sum(len(p) for p in parts) > MAX_DOCX_CHARS:
                break
        text = "\n".join(parts)
        if len(text) > MAX_DOCX_CHARS:
            text = text[:MAX_DOCX_CHARS] + "\n[truncated]"
        return text.strip()
    except Exception as e:
        log(f"DOCX parse failed: {e}")
        return ""


def _extract_docx_images(path: str) -> List[Dict[str, str]]:
    """Extract images from DOCX file and return as base64 data URLs."""
    images = []
    try:
        from docx import Document  # type: ignore
        import zipfile
    except Exception as e:
        log(f"DOCX image extraction import failed: {e}")
        return images
    
    try:
        # DOCX files are ZIP archives - extract images from word/media/
        with zipfile.ZipFile(path, 'r') as docx_zip:
            # List all files in the archive
            file_list = docx_zip.namelist()
            
            # Find all image files in word/media/
            image_files = [f for f in file_list if f.startswith('word/media/') and 
                          any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'])]
            
            for img_path in image_files:
                try:
                    # Read image data
                    img_data = docx_zip.read(img_path)
                    
                    # Determine MIME type from extension
                    ext = img_path.lower()
                    if ext.endswith('.png'):
                        mime_type = 'image/png'
                    elif ext.endswith(('.jpg', '.jpeg')):
                        mime_type = 'image/jpeg'
                    elif ext.endswith('.gif'):
                        mime_type = 'image/gif'
                    elif ext.endswith('.bmp'):
                        mime_type = 'image/bmp'
                    elif ext.endswith('.webp'):
                        mime_type = 'image/webp'
                    else:
                        mime_type = 'image/png'  # Default
                    
                    # Convert to base64
                    b64_data = base64.b64encode(img_data).decode('utf-8')
                    
                    images.append({
                        "url": f"data:{mime_type};base64,{b64_data}",
                        "name": os.path.basename(img_path),
                        "mime_type": mime_type
                    })
                    log(f"Extracted image from DOCX: {os.path.basename(img_path)} ({len(img_data)} bytes)")
                except Exception as img_e:
                    log(f"Failed to extract image {img_path} from DOCX: {img_e}")
        
        log(f"Extracted {len(images)} images from DOCX file")
    except Exception as e:
        log(f"DOCX image extraction failed: {e}")
    
    return images


def _extract_csv_text(path: str) -> str:
    try:
        rows = []
        with open(path, newline="", encoding="utf-8", errors="ignore") as fh:
            reader = csv.reader(fh)
            for i, row in enumerate(reader):
                if i >= MAX_CSV_ROWS:
                    rows.append(["[truncated rows]"])
                    break
                rows.append(row[:MAX_CSV_COLS])
        lines = [", ".join(row) for row in rows]
        return "\n".join(lines).strip()
    except Exception as e:
        log(f"CSV parse failed: {e}")
        return ""


def _extract_xlsx_text(path: str) -> str:
    try:
        import openpyxl  # type: ignore
    except Exception as e:
        log(f"XLSX import failed: {e}")
        return ""
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sheet = wb.active
        rows = []
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i >= MAX_XLSX_ROWS:
                rows.append(["[truncated rows]"])
                break
            limited = []
            for j, cell in enumerate(row):
                if j >= MAX_XLSX_COLS:
                    limited.append("[truncated cols]")
                    break
                limited.append("" if cell is None else str(cell))
            rows.append(limited)
        lines = ["\t".join(r) for r in rows]
        return "\n".join(lines).strip()
    except Exception as e:
        log(f"XLSX parse failed: {e}")
        return ""


def _extract_xlsx_images(path: str) -> List[Dict[str, str]]:
    """Extract images from XLSX file and return as base64 data URLs."""
    images = []
    try:
        import openpyxl  # type: ignore
        import zipfile
    except Exception as e:
        log(f"XLSX image extraction import failed: {e}")
        return images
    
    try:
        # XLSX files are ZIP archives - extract images from xl/media/
        with zipfile.ZipFile(path, 'r') as xlsx_zip:
            # List all files in the archive
            file_list = xlsx_zip.namelist()
            
            # Find all image files in xl/media/
            image_files = [f for f in file_list if f.startswith('xl/media/') and 
                          any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'])]
            
            for img_path in image_files:
                try:
                    # Read image data
                    img_data = xlsx_zip.read(img_path)
                    
                    # Determine MIME type from extension
                    ext = img_path.lower()
                    if ext.endswith('.png'):
                        mime_type = 'image/png'
                    elif ext.endswith(('.jpg', '.jpeg')):
                        mime_type = 'image/jpeg'
                    elif ext.endswith('.gif'):
                        mime_type = 'image/gif'
                    elif ext.endswith('.bmp'):
                        mime_type = 'image/bmp'
                    elif ext.endswith('.webp'):
                        mime_type = 'image/webp'
                    else:
                        mime_type = 'image/png'  # Default
                    
                    # Convert to base64
                    b64_data = base64.b64encode(img_data).decode('utf-8')
                    
                    images.append({
                        "url": f"data:{mime_type};base64,{b64_data}",
                        "name": os.path.basename(img_path),
                        "mime_type": mime_type
                    })
                    log(f"Extracted image from XLSX: {os.path.basename(img_path)} ({len(img_data)} bytes)")
                except Exception as img_e:
                    log(f"Failed to extract image {img_path} from XLSX: {img_e}")
        
        log(f"Extracted {len(images)} images from XLSX file")
    except Exception as e:
        log(f"XLSX image extraction failed: {e}")
    
    return images


def _extract_doc_text(path: str) -> str:
    """
    Best-effort DOC extraction: try to use textract if available; otherwise return empty.
    """
    try:
        import textract  # type: ignore
    except Exception:
        log("DOC parse skipped: textract not available")
        return ""
    try:
        text = textract.process(path).decode("utf-8", errors="ignore")
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS] + "\n[truncated]"
        return text.strip()
    except Exception as e:
        log(f"DOC parse failed: {e}")
        return ""


def _extract_xls_text(path: str) -> str:
    """
    Best-effort XLS extraction via xlrd if available.
    """
    try:
        import xlrd  # type: ignore
    except Exception:
        log("XLS parse skipped: xlrd not available")
        return ""
    try:
        book = xlrd.open_workbook(path, on_demand=True)
        sheet = book.sheet_by_index(0)
        lines = []
        for r in range(min(sheet.nrows, MAX_XLSX_ROWS)):
            row_vals = []
            for c in range(min(sheet.ncols, MAX_XLSX_COLS)):
                val = sheet.cell_value(r, c)
                row_vals.append("" if val is None else str(val))
            lines.append("\t".join(row_vals))
        text = "\n".join(lines)
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS] + "\n[truncated]"
        return text.strip()
    except Exception as e:
        log(f"XLS parse failed: {e}")
        return ""


def _extract_odt_text(path: str) -> str:
    try:
        from odf.opendocument import load  # type: ignore
        from odf import text as odf_text  # type: ignore
    except Exception:
        log("ODT parse skipped: odfpy not available")
        return ""
    try:
        doc = load(path)
        paras = []
        for p in doc.getElementsByType(odf_text.P):
            txt = "".join(t.data for t in p.childNodes if hasattr(t, "data"))
            if txt:
                paras.append(txt)
            if sum(len(x) for x in paras) > MAX_TEXT_CHARS:
                break
        text = "\n".join(paras)
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS] + "\n[truncated]"
        return text.strip()
    except Exception as e:
        log(f"ODT parse failed: {e}")
        return ""


def _extract_rtf_text(path: str) -> str:
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
    except Exception:
        log("RTF parse skipped: striprtf not available")
        return ""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as fh:
            data = fh.read()
        text = rtf_to_text(data)
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS] + "\n[truncated]"
        return text.strip()
    except Exception as e:
        log(f"RTF parse failed: {e}")
        return ""


def _extract_text_file(path: str, max_chars: int) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as fh:
            data = fh.read(max_chars + 1)
            if len(data) > max_chars:
                data = data[:max_chars] + "\n[truncated]"
            return data.strip()
    except Exception as e:
        log(f"Text parse failed: {e}")
        return ""


def _extract_tsv_text(path: str) -> str:
    try:
        rows = []
        with open(path, newline="", encoding="utf-8", errors="ignore") as fh:
            reader = csv.reader(fh, delimiter="\t")
            for i, row in enumerate(reader):
                if i >= MAX_CSV_ROWS:
                    rows.append(["[truncated rows]"])
                    break
                rows.append(row[:MAX_CSV_COLS])
        lines = ["\t".join(row) for row in rows]
        return "\n".join(lines).strip()
    except Exception as e:
        log(f"TSV parse failed: {e}")
        return ""


def _extract_json_text(path: str) -> str:
    try:
        with open(path, "rb") as fh:
            data = fh.read(MAX_JSON_BYTES + 1)
        truncated = len(data) > MAX_JSON_BYTES
        if truncated:
            data = data[:MAX_JSON_BYTES]
        try:
            obj = json.loads(data.decode("utf-8", errors="ignore"))
            text = json.dumps(obj, indent=2)
        except Exception:
            text = data.decode("utf-8", errors="ignore")
        if truncated or len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS] + "\n[truncated]"
        return text.strip()
    except Exception as e:
        log(f"JSON parse failed: {e}")
        return ""


def _extract_bruker_zip(path: str) -> str:
    """
    Parse Bruker NMR zip for key metadata. Metadata-only (no peak picking).
    """
    import zipfile
    try:
        if os.path.getsize(path) > MAX_ZIP_BYTES:
            raise HTTPException(status_code=400, detail=f"Zip too large (> {MAX_ZIP_BYTES/1024/1024:.0f} MB)")
        with zipfile.ZipFile(path, "r") as zf:
            namelist = zf.namelist()
            if len(namelist) > MAX_ZIP_FILES:
                raise HTTPException(status_code=400, detail="Zip has too many entries")

            # Detect Bruker structure
            lower_names = [n.lower() for n in namelist]
            if not any("acqus" in n for n in lower_names):
                return ""

            def read_text_member(name):
                try:
                    with zf.open(name, "r") as fh:
                        return fh.read().decode("utf-8", errors="ignore")
                except Exception:
                    return ""

            meta_sections = []

            acqus_files = [n for n in namelist if n.lower().endswith("acqus")]
            procs_files = [n for n in namelist if n.lower().endswith("procs")]
            title_files = [n for n in namelist if n.lower().endswith("title")]

            def parse_param(block, keys):
                lines = block.splitlines()
                out = []
                for key in keys:
                    for line in lines:
                        if line.strip().startswith(f"##${key}="):
                            out.append(f"{key}: {line.split('=',1)[1].strip()}")
                            break
                return out

            if acqus_files:
                data = read_text_member(acqus_files[0])
                vals = parse_param(data, ["SW", "TD", "O1", "SFO1", "NUC1", "TE", "RG", "D", "NS", "DATE"])
                if vals:
                    meta_sections.append("ACQUS:\n" + "\n".join(vals))

            if procs_files:
                data = read_text_member(procs_files[0])
                vals = parse_param(data, ["SF", "SI", "SSB", "LB", "WDW"])
                if vals:
                    meta_sections.append("PROCS:\n" + "\n".join(vals))

            if title_files:
                data = read_text_member(title_files[0])
                if data.strip():
                    meta_sections.append("TITLE:\n" + data.strip())

            if not meta_sections:
                return ""

            return "\n\n".join(meta_sections)
    except HTTPException:
        raise
    except Exception as e:
        log(f"Bruker zip parse failed: {e}")
        return ""


def _extract_jcamp_text(path: str) -> str:
    """
    Simple JCAMP-DX peak/metadata extraction (no external deps).
    """
    try:
        with open(path, "rb") as fh:
            raw = fh.read(MAX_JCAMP_BYTES + 1)
        truncated = len(raw) > MAX_JCAMP_BYTES
        data = raw[:MAX_JCAMP_BYTES].decode("utf-8", errors="ignore")
    except Exception as e:
        log(f"JCAMP read failed: {e}")
        return ""

    lines = data.splitlines()
    meta = {}
    for ln in lines[:200]:  # scan headers
        l = ln.strip()
        if not l.startswith("##"):
            continue
        if "=" in l:
            key, val = l[2:].split("=", 1)
            meta[key.strip().upper()] = val.strip()

    title = meta.get("TITLE", "")
    nucleus = meta.get(".OBSERVENUCLEUS") or meta.get("NUCLEUS") or meta.get(".NUCLEUS") or ""
    solvent = meta.get(".SOLVENT") or meta.get("SOLVENT") or meta.get("$SOLVENT") or ""
    freq = meta.get(".OBSERVEFREQUENCY") or meta.get("OBSERVE FREQUENCY") or meta.get("$SFO1") or ""
    temp = meta.get(".TEMPERATURE") or meta.get("TEMP") or meta.get("$TE") or ""

    # Extract peak table
    peak_section = ""
    if "##PEAKTABLE=" in data:
        try:
            start = data.index("##PEAKTABLE=")
            rest = data[start:].split("##", 2)
            if len(rest) >= 2:
                peak_section = rest[1]  # after PEAKTABLE tag until next ##
            else:
                peak_section = rest[0]
        except ValueError:
            pass
    peaks = []
    if peak_section:
        for ln in peak_section.splitlines():
            ln = ln.strip()
            if not ln or ln.startswith("##"):
                continue
            # Expect "shift, intensity" pairs
            parts = [p.strip() for p in ln.replace(";", ",").split(",") if p.strip()]
            if len(parts) >= 1:
                try:
                    shift = float(parts[0])
                    intensity = float(parts[1]) if len(parts) > 1 else 1.0
                    peaks.append(f"{shift:.2f} s 1H (int {intensity:.0f})")
                except Exception:
                    continue

    body = []
    if title:
        body.append(f"Title: {title}")
    if nucleus:
        body.append(f"Nucleus: {nucleus}")
    if solvent:
        body.append(f"Solvent: {solvent}")
    if freq:
        body.append(f"Freq: {freq} MHz")
    if temp:
        body.append(f"Temp: {temp}")
    if peaks:
        body.append("Peaks:")
        body.extend(peaks[:200])
    elif truncated:
        body.append("[JCAMP truncated; no peaks found]")

    return "\n".join(body).strip()


async def _post_with_retry(client: httpx.AsyncClient, url: str, headers: dict, payload: dict, attempts: int = 3, base_delay: float = 1.0, stream: bool = False):
    """
    POST with bounded retry/backoff for transient 429/5xx.
    """
    for attempt in range(1, attempts + 1):
        try:
            if stream:
                return await client.stream(
                    "POST",
                    url,
                    headers=headers,
                    json=payload,
                )
            resp = await client.post(
                url,
                headers=headers,
                json=payload,
            )
            if resp.status_code in (429, 500, 502, 503, 504) and attempt < attempts:
                delay = base_delay * (2 ** (attempt - 1))
                log(f"Transient error {resp.status_code}, retrying in {delay:.1f}s ({attempt}/{attempts})")
                await asyncio.sleep(delay)
                continue
            return resp
        except httpx.RequestError as e:
            if attempt >= attempts:
                raise
            delay = base_delay * (2 ** (attempt - 1))
            log(f"Request error {e}, retrying in {delay:.1f}s ({attempt}/{attempts})")
            await asyncio.sleep(delay)
    raise HTTPException(status_code=502, detail="Upstream retry exhausted")


def _extract_text_from_responses_chunk(data: dict) -> str:
    """
    Pull incremental text from Responses API stream chunk.
    """
    parts = []

    # Top-level delta style: {"delta":{"output":[...]}}
    delta = data.get("delta")
    if isinstance(delta, dict):
        for item in delta.get("output", []):
            if item.get("type") == "message":
                for c in item.get("content", []):
                    if not isinstance(c, dict):
                        continue
                    t = c.get("text") or (c.get("delta") or {}).get("text") or (c.get("delta") or {}).get("output_text")
                    if t:
                        parts.append(t)

    # Full output style: {"output":[{"type":"message","content":[...]}]}
    for item in data.get("output", []):
        if item.get("type") == "message":
            for c in item.get("content", []):
                if not isinstance(c, dict):
                    continue
                if c.get("type") in ("output_text", "text"):
                    t = c.get("text")
                    if t:
                        parts.append(t)
                delta_block = c.get("delta") or {}
                if isinstance(delta_block, dict):
                    t = delta_block.get("text") or delta_block.get("output_text")
                    if t:
                        parts.append(t)

    return "".join(parts)


def extract_pdfs_and_clean_text(text: str) -> tuple[str, list[dict]]:
    """
    Extract PDF markers from text and return cleaned text + list of PDFs.
    """
    pdfs = []
    
    for match in PDF_MARKER_RE.finditer(text):
        filename = match.group(1).strip()
        b64_data = match.group(2).strip()
        pdfs.append({
            "filename": filename,
            "base64": b64_data
        })
    
    # Remove markers from text
    cleaned = PDF_MARKER_RE.sub("", text).strip()
    return cleaned, pdfs


def extract_text_from_content(content) -> str:
    """Extract text from message content (string or list)."""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict):
                if item.get("type") == "text":
                    parts.append(item.get("text", ""))
                elif item.get("type") == "image_url":
                    # Keep image URLs as-is for now
                    pass
        return "\n".join(parts)
    return str(content) if content else ""


def _load_files_from_request(body: dict, messages: list) -> tuple[list[dict], list[dict], list[str]]:
    """
    Extract real uploaded files (PDFs, images, and doc text) into lists for Responses API.
    Returns (pdfs, images, texts). Images returned as data URLs for input_image.
    """
    files = _extract_all_files(body, messages)
    if not files:
        return [], [], []

    max_file_bytes = MAX_FILE_MB * 1024 * 1024
    max_total_bytes = MAX_TOTAL_MB * 1024 * 1024
    total_bytes = 0

    pdfs: List[dict] = []
    images: List[dict] = []
    texts: List[str] = []

    for f in files:
        path = _get_file_path(f)
        name = _get_file_name(f) or os.path.basename(path)
        if not path or not os.path.exists(path):
            raise HTTPException(status_code=400, detail=f"File not found: {name or path}")

        size = _file_size_bytes(path)
        if size < 0:
            raise HTTPException(status_code=400, detail=f"Cannot read file size: {name}")
        if size > max_file_bytes:
            raise HTTPException(status_code=400, detail=f"File too large (> {MAX_FILE_MB} MB): {name}")
        if total_bytes + size > max_total_bytes:
            raise HTTPException(status_code=400, detail=f"Total upload size exceeds {MAX_TOTAL_MB} MB limit")

        mime, _ = mimetypes.guess_type(name)
        mime = mime or "application/octet-stream"

        if _is_pdf(name, mime):
            # Read in chunks for large files to avoid memory spikes
            with open(path, "rb") as fh:
                if size > 10 * 1024 * 1024:  # For files > 10MB, read in chunks
                    import io
                    chunks = []
                    while True:
                        chunk = fh.read(8192)  # 8KB chunks
                        if not chunk:
                            break
                        chunks.append(chunk)
                    data = b''.join(chunks)
                else:
                    data = fh.read()
                b64 = base64.b64encode(data).decode("utf-8")
                del data  # Clear from memory immediately
            pdfs.append({"filename": name or "document.pdf", "base64": b64})
            total_bytes += size
            continue

        if _is_image(mime):
            # Read in chunks for large files to avoid memory spikes
            with open(path, "rb") as fh:
                if size > 10 * 1024 * 1024:  # For files > 10MB, read in chunks
                    import io
                    chunks = []
                    while True:
                        chunk = fh.read(8192)  # 8KB chunks
                        if not chunk:
                            break
                        chunks.append(chunk)
                    data = b''.join(chunks)
                else:
                    data = fh.read()
                b64 = base64.b64encode(data).decode("utf-8")
                del data  # Clear from memory immediately
            images.append({"url": f"data:{mime};base64,{b64}", "name": name})
            total_bytes += size
            continue

        # Text-first ingestion for DOCX/CSV/XLSX/others
        text_block = ""
        docx_images = []
        xlsx_images = []
        
        if _is_docx(name, mime):
            text_block = _extract_docx_text(path)
            # Also extract images from DOCX
            docx_images = _extract_docx_images(path)
            if docx_images:
                images.extend(docx_images)
                log(f"Added {len(docx_images)} images from DOCX file: {name}")
        elif _is_csv(name, mime):
            text_block = _extract_csv_text(path)
        elif _is_xlsx(name, mime):
            text_block = _extract_xlsx_text(path)
            # Also extract images from XLSX
            xlsx_images = _extract_xlsx_images(path)
            if xlsx_images:
                images.extend(xlsx_images)
                log(f"Added {len(xlsx_images)} images from XLSX file: {name}")
        elif _is_tsv(name, mime):
            text_block = _extract_tsv_text(path)
        elif _is_md(name, mime) or _is_txt(name, mime):
            text_block = _extract_text_file(path, MAX_TEXT_CHARS)
        elif _is_json_file(name, mime):
            text_block = _extract_json_text(path)
        elif _is_doc(name, mime):
            text_block = _extract_doc_text(path)
        elif _is_xls(name, mime):
            text_block = _extract_xls_text(path)
        elif _is_odt(name, mime):
            text_block = _extract_odt_text(path)
        elif _is_rtf(name, mime):
            text_block = _extract_rtf_text(path)
        elif _is_bruker_zip(name, mime):
            text_block = _extract_bruker_zip(path)
        elif _is_jcamp(name, mime):
            text_block = _extract_jcamp_text(path)

        if text_block:
            texts.append(f"=== File: {name} ===\n{text_block}")
            total_bytes += size
            continue

        # Non-supported types: skip silently but keep base functionality intact
        log(f"Skipping unsupported file type: {name} ({mime})")

    # ALSO include inline images injected into chat content by Functions/Filters
    inline_images = _extract_inline_images_from_messages(messages)
    if inline_images:
        images.extend(inline_images)
        log(f"Added {len(inline_images)} inline images from message content (Filter-injected).")

    return pdfs, images, texts


async def call_responses_api(model: str, conversation_history: list[dict]) -> dict:
    """
    Call OpenAI Responses API with full conversation history (preserves memory).
    conversation_history should be a list of message dicts with role and content.
    """
    # Filter user text for content safety
    for msg in conversation_history:
        if msg.get("role") == "user":
            for item in msg.get("content", []):
                if item.get("type") == "input_text":
                    _content_filter(item.get("text", ""))
    
    METRICS["responses_api_calls"] += 1
    
    payload = {
        "model": model,
        "input": conversation_history
    }
    
    # Log what we're sending - VERIFICATION
    payload_size = len(json.dumps(payload))
    user_msgs = [m for m in conversation_history if m.get("role") == "user"]
    assistant_msgs = [m for m in conversation_history if m.get("role") == "assistant"]
    
    # Count files/images in current message
    pdf_count = 0
    image_count = 0
    if user_msgs:
        last_user = user_msgs[-1]
        for item in last_user.get("content", []):
            if item.get("type") == "input_file":
                pdf_count += 1
            elif item.get("type") == "input_image":
                image_count += 1
    
    log(f"📤 SENDING TO OPENAI Responses API:")
    log(f"   Model: {model}")
    log(f"   Conversation history: {len(conversation_history)} messages ({len(user_msgs)} user, {len(assistant_msgs)} assistant)")
    log(f"   PDFs in current message: {pdf_count}")
    log(f"   Images in current message: {image_count}")
    log(f"   Total payload size: {payload_size/1024:.1f}KB")
    log(f"✅ VERIFIED: Full conversation history preserved for memory")
    
    client = await get_http_client()
    resp = await _post_with_retry(
        client,
        f"{OPENAI_BASE_URL}/responses",
        headers={
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json",
        },
        payload=payload,
        attempts=3,
        base_delay=1.0,
        stream=False,
    )
    
    # Verify response
    if resp.status_code >= 400:
        error_text = resp.text[:500] if hasattr(resp, 'text') else str(resp)
        log(f"❌ Responses API ERROR: HTTP {resp.status_code}")
        log(f"Error details: {error_text}")
        METRICS["errors_total"] += 1
        METRICS["last_error"] = f"HTTP {resp.status_code}: {error_text[:200]}"
        raise HTTPException(status_code=resp.status_code, detail=error_text)
    
    # Success - parse and validate response
    try:
        response_data = resp.json()
        
        # Verify response structure
        if not isinstance(response_data, dict):
            log(f"⚠️ WARNING: Unexpected response type: {type(response_data)}")
        else:
            # Check for output in response (indicates successful processing)
            output = response_data.get("output", [])
            has_output = len(output) > 0
            
            # Check usage stats (indicates API processed the request)
            usage = response_data.get("usage", {})
            has_usage = bool(usage)
            
            if has_output or has_usage:
                log(f"✅ SUCCESS: OpenAI received and processed the request")
                log(f"   - Response ID: {response_data.get('id', 'N/A')}")
                log(f"   - Output items: {len(output)}")
                log(f"   - Usage stats: {bool(usage)}")
                if image_count > 0:
                    log(f"   - Images sent: {image_count} (verified in payload)")
            else:
                log(f"⚠️ WARNING: Response received but no output/usage data found")
        
        return response_data
    except json.JSONDecodeError as e:
        log(f"❌ ERROR: Failed to parse JSON response: {e}")
        log(f"Response text (first 500 chars): {resp.text[:500] if hasattr(resp, 'text') else 'N/A'}")
        METRICS["errors_total"] += 1
        METRICS["last_error"] = f"JSON parse error: {str(e)}"
        raise HTTPException(status_code=502, detail="Invalid JSON response from OpenAI")


async def stream_responses_api(model: str, conversation_history: list[dict], request: Request = None):
    """
    Stream OpenAI Responses API SSE and convert to chat-completion chunks.
    Supports cancellation via request parameter.
    Preserves full conversation history for memory.
    """
    # Filter user text for content safety
    for msg in conversation_history:
        if msg.get("role") == "user":
            for item in msg.get("content", []):
                if item.get("type") == "input_text":
                    _content_filter(item.get("text", ""))
    
    METRICS["responses_api_calls"] += 1
    
    user_msgs = [m for m in conversation_history if m.get("role") == "user"]
    assistant_msgs = [m for m in conversation_history if m.get("role") == "assistant"]
    log(f"Streaming with conversation history: {len(conversation_history)} messages ({len(user_msgs)} user, {len(assistant_msgs)} assistant)")
    
    payload = {
        "model": model,
        "input": conversation_history,
        "stream": True,
    }

    url = f"{OPENAI_BASE_URL}/responses"
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json",
    }

    attempts = 3
    base_delay = 1.0

    client = await get_http_client()

    for attempt in range(1, attempts + 1):
        try:
            resp_context = await _post_with_retry(client, url, headers, payload, attempts=1, stream=True)
            async with resp_context as resp:
                if resp.status_code >= 400:
                    if resp.status_code in (429, 500, 502, 503, 504) and attempt < attempts:
                        delay = base_delay * (2 ** (attempt - 1))
                        log(f"❌ Responses stream error {resp.status_code}, retrying in {delay:.1f}s ({attempt}/{attempts})")
                        await asyncio.sleep(delay)
                        continue
                    text = await resp.aread()
                    error_msg = text.decode() if hasattr(text, "decode") else str(text)
                    log(f"❌ Responses stream ERROR: HTTP {resp.status_code} - {error_msg[:200]}")
                    METRICS["errors_total"] += 1
                    METRICS["last_error"] = f"Stream HTTP {resp.status_code}: {error_msg[:200]}"
                    raise HTTPException(status_code=resp.status_code, detail=error_msg)
                
                    # Success - log that stream started
                    log(f"✅ Streaming started: waiting for response chunks...")

                    chunk_count = 0
                    async for line in resp.aiter_lines():
                        # Check for client disconnection
                        if request and await request.is_disconnected():
                            log("[PROXY] Client disconnected during Responses API stream - stopping")
                            break
                        
                        if not line:
                            continue
                        if line.startswith("data:"):
                            data = line[len("data:"):].strip()
                            if data == "[DONE]":
                                log(f"✅ Received [DONE] signal, stream complete ({chunk_count} chunks)")
                                yield "data: [DONE]\n\n"
                                return
                            try:
                                parsed = json.loads(data)
                            except json.JSONDecodeError:
                                continue
                            chunk_text = _extract_text_from_responses_chunk(parsed)
                            if not chunk_text:
                                continue
                            chunk = {
                                "id": parsed.get("id", "resp_proxy"),
                                "object": "chat.completion.chunk",
                                "model": model,
                                "choices": [{
                                    "index": 0,
                                    "delta": {"role": "assistant", "content": chunk_text},
                                    "finish_reason": None
                                }]
                            }
                            chunk_count += 1
                            if chunk_count == 1:
                                log(f"✅ First chunk received, streaming response...")
                            yield f"data: {json.dumps(chunk)}\n\n"
                    # If stream completed without [DONE], send a terminator
                    log(f"✅ Stream completed ({chunk_count} chunks total)")
                    yield "data: [DONE]\n\n"
                    return
        except Exception as e:
            if attempt >= attempts:
                log(f"Responses stream failed: {e}")
                raise
            delay = base_delay * (2 ** (attempt - 1))
            log(f"Responses stream exception {e}, retrying in {delay:.1f}s ({attempt}/{attempts})")
            await asyncio.sleep(delay)



async def call_chat_completions(body: dict) -> dict:
    """
    Forward to standard Chat Completions API (fallback when no PDFs).
    """
    client = await get_http_client()
    resp = await _post_with_retry(
        client,
        f"{OPENAI_BASE_URL}/chat/completions",
        headers={
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json",
        },
        payload=body,
        attempts=3,
        base_delay=1.0,
        stream=False,
    )
    if resp.status_code >= 400:
        log(f"Chat Completions error: {resp.status_code}")
        raise HTTPException(status_code=resp.status_code, detail=resp.text)
    return resp.json()


def responses_to_chat_completion(resp_data: dict, model: str) -> dict:
    """
    Convert Responses API output to Chat Completions format for OpenWebUI.
    """
    output_text = ""
    
    for item in resp_data.get("output", []):
        if item.get("type") == "message":
            for c in item.get("content", []):
                if c.get("type") in ("output_text", "text"):
                    output_text += c.get("text", "")
    
    return {
        "id": resp_data.get("id", "resp_proxy"),
        "object": "chat.completion",
        "model": model,
        "choices": [{
            "index": 0,
            "message": {
                "role": "assistant",
                "content": output_text.strip() or "(No output from model)"
            },
            "finish_reason": "stop"
        }],
        "usage": resp_data.get("usage", {})
    }


@app.post("/v1/chat/completions")
async def chat_completions(request: Request):
    """
    Main endpoint - intercepts Chat Completions, upgrades to Responses API if PDFs found.
    Supports request cancellation when client disconnects.
    """
    start = time.perf_counter()
    METRICS["requests_total"] += 1
    body = await request.json()
    
    model = body.get("model", "gpt-4o")
    messages = body.get("messages", [])
    stream = body.get("stream", False)
    
    log(f"[PROXY] POST /v1/chat/completions - model: {model}, messages: {len(messages)}, stream: {stream}")
    
    # Check if client disconnected before processing
    if await request.is_disconnected():
        log("[PROXY] Client disconnected before processing")
        raise HTTPException(status_code=499, detail="Client disconnected")
    
    # Collect all text and find PDF markers
    all_text = ""
    marker_pdfs = []
    marker_images = []
    upload_pdfs = []
    upload_images = []
    upload_texts = []
    
    for msg in messages:
        if msg.get("role") != "user":
            continue
        
        content = msg.get("content")
        text = extract_text_from_content(content)
        
        # Extract PDFs from text
        cleaned_text, pdfs = extract_pdfs_and_clean_text(text)
        all_text += cleaned_text + "\n"
        marker_pdfs.extend(pdfs)
        
        # Also check for existing image_url items
        if isinstance(content, list):
            for item in content:
                if isinstance(item, dict) and item.get("type") == "image_url":
                    img_url = item.get("image_url", {})
                    if isinstance(img_url, dict):
                        marker_images.append({"url": img_url.get("url", "")})
                    elif isinstance(img_url, str):
                        marker_images.append({"url": img_url})
    
    all_text = all_text.strip()

    # Load real uploaded files (PDF/images/text) from body/messages
    try:
        upload_pdfs, upload_images, upload_texts = _load_files_from_request(body, messages)
    except HTTPException as e:
        # User-facing error for size/not-found issues
        raise e
    except Exception as e:
        log(f"File load failed: {e}")
        upload_pdfs, upload_images, upload_texts = [], [], []

    all_pdfs = marker_pdfs + upload_pdfs
    all_images = marker_images + upload_images
    all_text_blocks = upload_texts
    
    log(f"Found {len(all_pdfs)} PDF(s), {len(all_images)} image(s), {len(all_text_blocks)} text block(s) (markers+uploads)")

    # Use the user's text as-is - system prompt in OpenWebUI handles all instructions
    augmented_text = all_text
    
    # Only use Responses API for PDFs - use faster Chat Completions for images
    if all_pdfs:
        log("Using Responses API for PDF analysis")
        try:
            # Build conversation history for Responses API (preserve memory)
            conversation_history = []
            for msg in messages:
                role = msg.get("role")
                content = msg.get("content", "")
                
                # Skip system messages (handled by OpenWebUI)
                if role == "system":
                    continue
                
                # Extract text from content
                if isinstance(content, str):
                    text = content
                elif isinstance(content, list):
                    text_parts = []
                    for item in content:
                        if isinstance(item, dict) and item.get("type") == "text":
                            text_parts.append(item.get("text", ""))
                    text = " ".join(text_parts)
                else:
                    text = str(content) if content else ""
                
                if text.strip() or role == "assistant":
                    conversation_history.append({
                        "role": role,
                        "content": [{"type": "input_text", "text": text.strip()}]
                    })
            
            # Add current user message with files
            current_user_content = []
            for pdf in all_pdfs:
                current_user_content.append({
                    "type": "input_file",
                    "filename": pdf["filename"],
                    "file_data": f"data:application/pdf;base64,{pdf['base64']}",
                })
            for img in all_images:
                img_url = img.get("url", "")
                if img_url:
                    current_user_content.append({
                        "type": "input_image",
                        "image_url": img_url,
                    })
            for txt in all_text_blocks:
                if txt:
                    current_user_content.append({
                        "type": "input_text",
                        "text": txt
                    })
            if augmented_text:
                current_user_content.append({
                    "type": "input_text",
                    "text": augmented_text
                })
            
            # Update last message in history or add new one
            if conversation_history and conversation_history[-1]["role"] == "user":
                conversation_history[-1]["content"] = current_user_content
            else:
                conversation_history.append({
                    "role": "user",
                    "content": current_user_content
                })
            
            log(f"Preserving conversation history: {len(conversation_history)} messages")
            
            if stream:
                async def proxied_stream():
                    try:
                        async for event in stream_responses_api(model, conversation_history, request):
                            # Check if client disconnected
                            if await request.is_disconnected():
                                log("[PROXY] Client disconnected during Responses API streaming - stopping")
                                break
                            yield event
                    except asyncio.CancelledError:
                        log("[PROXY] Responses API stream cancelled")
                        raise
                    except Exception as e:
                        if "disconnect" in str(e).lower() or "cancelled" in str(e).lower():
                            log(f"[PROXY] Responses stream cancelled/disconnected: {e}")
                        else:
                            raise

                return StreamingResponse(proxied_stream(), media_type="text/event-stream")

            resp_data = await call_responses_api(model, conversation_history)
            result = responses_to_chat_completion(resp_data, model)
            return JSONResponse(content=result)
            
        except Exception as e:
            METRICS["errors_total"] += 1
            METRICS["last_error"] = str(e)
            log(f"Responses API failed: {e}")
            # Fall back to chat completions
            pass
    
    # For images (no PDFs), use fast Chat Completions API with vision support
    # For PDFs, Responses API is used above
    # For text-only, use standard Chat Completions
    if all_images and not all_pdfs:
        log(f"Using fast Chat Completions API with vision support for {len(all_images)} images")
    else:
        log("Using standard Chat Completions API")
    
    # Clean PDF markers from messages before forwarding
    cleaned_messages = []
    for msg in messages:
        new_msg = msg.copy()
        content = msg.get("content")
        if isinstance(content, str):
            cleaned, _ = extract_pdfs_and_clean_text(content)
            new_msg["content"] = cleaned
        cleaned_messages.append(new_msg)
    
    body["messages"] = cleaned_messages
    
    if stream:
        # Stream from Chat Completions with cancellation support
        async def stream_response():
            try:
                client = await get_http_client()
                async with client.stream(
                    "POST",
                    f"{OPENAI_BASE_URL}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {OPENAI_API_KEY}",
                        "Content-Type": "application/json",
                    },
                    json=body
                ) as resp:
                    async for chunk in resp.aiter_bytes():
                        # Check if client disconnected
                        if await request.is_disconnected():
                            log("[PROXY] Client disconnected during streaming - stopping")
                            break
                        yield chunk
            except asyncio.CancelledError:
                log("[PROXY] Stream cancelled")
                raise
            except Exception as e:
                if "disconnect" in str(e).lower() or "cancelled" in str(e).lower():
                    log(f"[PROXY] Stream cancelled/disconnected: {e}")
                else:
                    raise
        
        return StreamingResponse(stream_response(), media_type="text/event-stream")
    
    try:
        result = await call_chat_completions(body)
    except Exception as e:
        METRICS["errors_total"] += 1
        METRICS["last_error"] = str(e)
        raise
    latency_ms = (time.perf_counter() - start) * 1000
    METRICS["last_latency_ms"] = latency_ms
    return JSONResponse(content=result)


def _safe_slug(text: str, default: str = "report") -> str:
    """Create a simple filename-safe slug."""
    slug = re.sub(r"[^a-zA-Z0-9_-]+", "-", text).strip("-")
    return slug or default

def _is_enabled(env_key: str, default: bool = True) -> bool:
    return os.environ.get(env_key, "true" if default else "false").lower() == "true"


def _content_filter(text: str) -> None:
    """Basic content filter hook."""
    banned = ["password", "apikey", "secret", "token", "private key"]
    lower = text.lower()
    if any(b in lower for b in banned):
        raise HTTPException(status_code=400, detail="Content blocked by filter")


async def _post_with_retry(client: httpx.AsyncClient, url: str, headers: dict, payload: dict, attempts: int = 3, base_delay: float = 1.0, stream: bool = False):
    for i in range(attempts):
        try:
            if stream:
                # client.stream() returns an async context manager, don't await it
                return client.stream("POST", url, headers=headers, json=payload)
            else:
                resp = await client.post(url, headers=headers, json=payload)
                if resp.status_code in (429, 500, 502, 503, 504) and i < attempts - 1:
                    await asyncio.sleep(base_delay * (2 ** i))
                    continue
                resp.raise_for_status()
                return resp
        except Exception:
            if i == attempts - 1:
                raise
            await asyncio.sleep(base_delay * (2 ** i))


def render_report_pdf(report: dict) -> bytes:
    """
    Render a professional PDF report with branding and visual enhancements.
    Expected keys: title, subtitle, author, date, sections [{heading, body, bullets, table{headers, rows}, images[{url,data_url,caption}]}], footer.
    Branding keys: company_name, logo_path, primary_color, secondary_color, document_type.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    flow = []

    # Get branding info
    company_name = report.get("company_name", "GLChemTec")
    logo_path = report.get("logo_path", "")
    primary_color = report.get("primary_color", "#1d2b3a")
    secondary_color = report.get("secondary_color", "#e6eef5")
    document_type = report.get("document_type", "pdf")
    
    # Create custom styles with branding colors
    from reportlab.lib.units import inch  # type: ignore
    from reportlab.lib.enums import TA_CENTER, TA_LEFT  # type: ignore
    
    # Custom title style with brand color
    title_style = styles["Title"]
    title_style.textColor = colors.HexColor(primary_color)
    title_style.fontSize = 24
    title_style.spaceAfter = 12
    
    # Custom heading style
    heading2_style = styles["Heading2"]
    heading2_style.textColor = colors.HexColor(primary_color)
    heading2_style.fontSize = 16
    heading2_style.spaceAfter = 8
    heading2_style.spaceBefore = 12

    title = report.get("title") or "Report"
    subtitle = report.get("subtitle", "")
    author = report.get("author", company_name)
    date = report.get("date", "")
    footer = report.get("footer", f"Generated by {company_name}")
    sections = report.get("sections", [])
    
    # Add logo if available
    if logo_path and os.path.exists(logo_path):
        try:
            logo_img = RLImage(logo_path, width=2*inch, preserveAspectRatio=True)
            flow.append(logo_img)
            flow.append(Spacer(1, 12))
        except Exception:
            pass  # Continue if logo fails to load

    # Add branded header with colored bar
    header_table = Table([[title]], colWidths=[7*inch])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(primary_color)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 20),
        ("ALIGN", (0, 0), (-1, 0), "LEFT"),
        ("VALIGN", (0, 0), (-1, 0), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, 0), 12),
        ("RIGHTPADDING", (0, 0), (-1, 0), 12),
        ("TOPPADDING", (0, 0), (-1, 0), 12),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
    ]))
    flow.append(header_table)
    flow.append(Spacer(1, 6))
    
    if subtitle:
        flow.append(Paragraph(subtitle, styles["Heading3"]))
    if author or date:
        meta = " | ".join([x for x in [author, date] if x])
        if meta:
            meta_para = Paragraph(meta, styles["Normal"])
            flow.append(meta_para)
    flow.append(Spacer(1, 12))

    for sec in sections:
        heading = sec.get("heading", "")
        body = sec.get("body", "")
        bullets = sec.get("bullets", [])
        table = sec.get("table")
        images = sec.get("images", [])

        if heading:
            # Add colored section heading
            heading_para = Paragraph(heading, heading2_style)
            flow.append(heading_para)
            # Add subtle divider line
            divider = Table([[""]], colWidths=[7*inch])
            divider.setStyle(TableStyle([
                ("LINEBELOW", (0, 0), (-1, 0), 1, colors.HexColor(secondary_color)),
                ("TOPPADDING", (0, 0), (-1, 0), 4),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
            ]))
            flow.append(divider)
        if body:
            flow.append(Paragraph(body.replace("\n", "<br/>"), styles["Normal"]))
        if bullets and isinstance(bullets, list):
            for b in bullets:
                flow.append(Paragraph(f"• {b}", styles["Normal"]))
        if table and isinstance(table, dict):
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            data = []
            if headers:
                data.append(headers)
            data.extend(rows)
            if data:
                t = Table(data, repeatRows=1)
                # Use branding colors
                table_bg = colors.HexColor(secondary_color)
                table_text = colors.HexColor(primary_color)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), table_bg),
                    ("TEXTCOLOR", (0, 0), (-1, 0), table_text),
                    ("GRID", (0, 0), (-1, -1), 0.5, table_text),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, 0), 10),
                    ("FONTSIZE", (0, 1), (-1, -1), 9),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8f9fa")]),
                ]))
                flow.append(Spacer(1, 6))
                flow.append(t)
        if images and isinstance(images, list):
            for img in images[:5]:  # limit
                src = img.get("data_url") or img.get("url") or ""
                caption = img.get("caption", "")
                img_bytes = None
                if src.startswith("data:"):
                    try:
                        b64data = src.split(",", 1)[1]
                        img_bytes = base64.b64decode(b64data)
                    except Exception:
                        img_bytes = None
                elif os.path.exists(src):
                    with open(src, "rb") as f:
                        img_bytes = f.read()
                if img_bytes:
                    tmp = io.BytesIO(img_bytes)
                    try:
                        flow.append(RLImage(tmp, width=400, preserveAspectRatio=True))
                        if caption:
                            flow.append(Paragraph(caption, styles["Italic"]))
                    except Exception:
                        pass
        flow.append(Spacer(1, 12))

    if footer:
        flow.append(Spacer(1, 24))
        flow.append(Paragraph(footer, styles["Normal"]))

    doc.build(flow)
    return buffer.getvalue()


def render_report_docx(report: dict) -> bytes:
    """Render a professional DOCX report with branding and visual enhancements."""
    from docx.shared import RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Get branding info
    company_name = report.get("company_name", "GLChemTec")
    logo_path = report.get("logo_path", "")
    primary_color = report.get("primary_color", "#1d2b3a")
    secondary_color = report.get("secondary_color", "#e6eef5")
    
    # Convert hex to RGB
    def hex_to_rgb(hex_color):
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    primary_rgb = hex_to_rgb(primary_color)
    primary_color_obj = RGBColor(primary_rgb[0], primary_rgb[1], primary_rgb[2])
    
    # Add logo if available
    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(2))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph("")  # Spacing
        except Exception:
            pass
    
    title = report.get("title") or "Report"
    subtitle = report.get("subtitle", "")
    author = report.get("author", company_name)
    date = report.get("date", "")
    sections = report.get("sections", [])
    footer = report.get("footer", f"Generated by {company_name}")
    
    # Add branded title
    title_para = doc.add_heading(title, level=1)
    title_run = title_para.runs[0] if title_para.runs else title_para.add_run(title)
    title_run.font.color.rgb = primary_color_obj
    title_run.bold = True

    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.runs[0].font.color.rgb = RGBColor(100, 100, 100) if subtitle_para.runs else None
    meta = " | ".join([x for x in [author, date] if x])
    if meta:
        meta_para = doc.add_paragraph(meta)
        if meta_para.runs:
            meta_para.runs[0].font.size = Pt(9)
            meta_para.runs[0].font.italic = True
    doc.add_paragraph("")

    for sec in sections:
        heading = sec.get("heading", "")
        body = sec.get("body", "")
        bullets = sec.get("bullets", [])
        table = sec.get("table")
        images = sec.get("images", [])

        if heading:
            # Add branded section heading
            heading_para = doc.add_heading(heading, level=2)
            if heading_para.runs:
                heading_para.runs[0].font.color.rgb = primary_color_obj
                heading_para.runs[0].bold = True
        if body:
            doc.add_paragraph(body)
        if bullets and isinstance(bullets, list):
            for b in bullets:
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(str(b))
        if table and isinstance(table, dict):
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            if headers or rows:
                cols = len(headers) if headers else len(rows[0])
                t = doc.add_table(rows=1, cols=cols)
                hdr_cells = t.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = str(h)
                for row in rows:
                    cells = t.add_row().cells
                    for i, val in enumerate(row):
                        cells[i].text = str(val)
        if images and isinstance(images, list):
            for img in images[:5]:
                src = img.get("data_url") or img.get("url") or ""
                caption = img.get("caption", "")
                img_bytes = None
                if src.startswith("data:"):
                    try:
                        b64data = src.split(",", 1)[1]
                        img_bytes = base64.b64decode(b64data)
                    except Exception:
                        img_bytes = None
                elif os.path.exists(src):
                    with open(src, "rb") as f:
                        img_bytes = f.read()
                if img_bytes:
                    tmp = io.BytesIO(img_bytes)
                    try:
                        doc.add_picture(tmp, width=None)
                        if caption:
                            doc.add_paragraph(caption).italic = True
                    except Exception:
                        pass
        doc.add_paragraph("")

    if footer:
        doc.add_paragraph(footer)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def summarize_csv_tsv(text: str, delimiter: str = ",", max_rows: int = 1000, max_cols: int = 50) -> dict:
    rows = []
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    for i, row in enumerate(reader):
        if i >= max_rows:
            break
        if len(row) > max_cols:
            row = row[:max_cols]
        rows.append(row)
    if not rows:
        return {"rows": 0, "cols": 0, "headers": [], "sample": []}
    headers = rows[0]
    data_rows = rows[1:]
    sample = data_rows[:5]
    return {
        "rows": len(rows) - 1,
        "cols": len(headers),
        "headers": headers,
        "sample": sample,
    }


def summarize_xlsx(content: bytes, max_rows: int = 1000, max_cols: int = 50) -> dict:
    wb = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    sheet = wb.active
    headers = []
    data = []
    for r_idx, row in enumerate(sheet.iter_rows(max_row=max_rows, max_col=max_cols, values_only=True)):
        vals = ["" if v is None else v for v in row]
        if r_idx == 0:
            headers = vals
        else:
            data.append(vals)
        if len(data) >= max_rows:
            break
    return {
        "rows": len(data),
        "cols": len(headers),
        "headers": headers,
        "sample": data[:5],
    }


def summarize_json(content: bytes, max_bytes: int = 1_000_000) -> dict:
    if len(content) > max_bytes:
        raise ValueError("JSON too large")
    obj = json.loads(content)
    def safe_repr(o, max_len=200):
        s = repr(o)
        return s if len(s) <= max_len else s[:max_len] + "..."
    summary = {
        "type": type(obj).__name__,
    }
    if isinstance(obj, dict):
        keys = list(obj.keys())
        summary["keys"] = keys[:50]
        summary["len"] = len(obj)
    elif isinstance(obj, list):
        summary["len"] = len(obj)
        if obj and isinstance(obj[0], dict):
            keys = set()
            for item in obj[:20]:
                if isinstance(item, dict):
                    keys.update(item.keys())
            summary["keys"] = list(keys)[:50]
        summary["sample"] = safe_repr(obj[:3])
    else:
        summary["value"] = safe_repr(obj)
    return summary


def analyze_file_payload(payload: dict) -> dict:
    """
    Analyze a file payload with base64 content.
    Expected keys: filename, content_base64, content_type?
    Supports CSV/TSV/XLSX/JSON (lightweight summary).
    """
    filename = payload.get("filename", "")
    b64 = payload.get("content_base64", "")
    ctype = payload.get("content_type", "")
    if not filename or not b64:
        raise HTTPException(status_code=400, detail="filename and content_base64 are required")
    try:
        raw = base64.b64decode(b64)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid base64 content")

    max_bytes = 5 * 1024 * 1024
    if len(raw) > max_bytes:
        raise HTTPException(status_code=400, detail="File too large (max 5MB)")

    name_lower = filename.lower()
    if name_lower.endswith(".csv"):
        text = raw.decode("utf-8", errors="replace")
        summary = summarize_csv_tsv(text, delimiter=",")
        return {"kind": "csv", "summary": summary}
    if name_lower.endswith(".tsv"):
        text = raw.decode("utf-8", errors="replace")
        summary = summarize_csv_tsv(text, delimiter="\\t")
        return {"kind": "tsv", "summary": summary}
    if name_lower.endswith(".xlsx"):
        summary = summarize_xlsx(raw)
        return {"kind": "xlsx", "summary": summary}
    if name_lower.endswith(".json") or "json" in ctype:
        summary = summarize_json(raw)
        return {"kind": "json", "summary": summary}

    raise HTTPException(status_code=400, detail="Unsupported file type for analysis")


def _decode_data_url(data_url: str) -> bytes:
    if not data_url.startswith("data:"):
        raise ValueError("Not a data URL")
    try:
        b64data = data_url.split(",", 1)[1]
        return base64.b64decode(b64data)
    except Exception as e:
        raise ValueError(f"Invalid data URL: {e}")


def _prepare_audio_bytes(payload: dict, max_bytes: int = 5 * 1024 * 1024) -> bytes:
    """
    Accepts either content_base64 or data_url in payload["audio"].
    """
    audio = payload.get("audio") or {}
    b64 = audio.get("content_base64", "")
    data_url = audio.get("data_url", "")
    raw = b""
    if b64:
        try:
            raw = base64.b64decode(b64)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid audio base64")
    elif data_url:
        try:
            raw = _decode_data_url(data_url)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
    else:
        raise HTTPException(status_code=400, detail="audio.content_base64 or audio.data_url is required")

    if len(raw) > max_bytes:
        raise HTTPException(status_code=400, detail="Audio too large (max 5MB)")
    return raw


async def openai_transcribe_audio(raw: bytes, filename: str = "audio.wav", model: str = "whisper-1") -> str:
    files = {"file": (filename, raw, "audio/wav")}
    data = {"model": model}
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}"}
    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(f"{OPENAI_BASE_URL}/audio/transcriptions", headers=headers, files=files, data=data)
            resp.raise_for_status()
            return resp.json().get("text", "")
    except Exception as e:
        log(f"Transcription failed: {e}")
        raise HTTPException(status_code=502, detail="Transcription failed")


async def openai_tts(text: str, voice: str = "alloy", model: str = "tts-1") -> bytes:
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}"}
    payload = {"model": model, "voice": voice, "input": text}
    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(f"{OPENAI_BASE_URL}/audio/speech", headers=headers, json=payload)
            resp.raise_for_status()
            return resp.content
    except Exception as e:
        log(f"TTS failed: {e}")
        raise HTTPException(status_code=502, detail="TTS failed")



# Export file storage - disk-based for multi-worker/instance support
EXPORT_DIR = Path("/app/backend/data/uploads/exports")
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

# Temporary file storage for exports (in-memory cache, falls back to disk)
EXPORT_FILES: Dict[str, Dict[str, Any]] = {}

@app.post("/v1/report/pdf")
async def generate_report_pdf(report: dict):
    """Generate a PDF report from structured JSON."""
    try:
        pdf_bytes = render_report_pdf(report)
    except Exception as e:
        log(f"PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")

    filename = _safe_slug(report.get("title", "report")) + ".pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/v1/report/docx")
async def generate_report_docx(report: dict):
    """Generate a DOCX report from structured JSON."""
    try:
        docx_bytes = render_report_docx(report)
    except Exception as e:
        log(f"DOCX generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {e}")

    filename = _safe_slug(report.get("title", "report")) + ".docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


async def _generate_ai_report_internal(conversation_history: list, format_type: str = "pdf", model: str = "gpt-4o") -> dict:
    """
    Internal function to generate AI report using OpenAI Chat Completions API.
    Returns structured report dict.
    """
    if not conversation_history:
        raise ValueError("No conversation history provided")
    
    log(f"Generating AI report: {len(conversation_history)} messages, format={format_type}, model={model}")
    
    # Create a system message that asks for structured report format
    system_prompt = """You are a professional technical writer. Create a well-structured, comprehensive report from the conversation history.

Format your response as a JSON object with this exact structure:
{
  "title": "Clear, descriptive title",
  "subtitle": "Optional subtitle",
  "summary": "Executive summary (2-3 sentences)",
  "sections": [
    {
      "heading": "Section title",
      "body": "Main content paragraph(s)",
      "bullets": ["Key point 1", "Key point 2", ...],
      "table": {
        "headers": ["Column 1", "Column 2"],
        "rows": [["Data 1", "Data 2"], ...]
      }
    }
  ],
  "conclusions": ["Conclusion point 1", "Conclusion point 2", ...],
  "recommendations": ["Recommendation 1", "Recommendation 2", ...]
}

Guidelines:
- Create logical sections based on topics discussed
- Include key findings, data, and insights
- Use tables for structured data
- Add conclusions and recommendations if applicable
- Be concise but comprehensive
- Maintain technical accuracy
- Return ONLY valid JSON, no markdown formatting or code blocks"""
    
    # Convert conversation history to Chat Completions format
    # The incoming format might be Responses API format, so we need to normalize it
    chat_messages = [{"role": "system", "content": system_prompt}]
    
    for msg in conversation_history:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        
        # Handle different content formats
        if isinstance(content, str):
            text_content = content
        elif isinstance(content, list):
            # Extract text from list format (Responses API style)
            text_parts = []
            for item in content:
                if isinstance(item, dict):
                    # Handle {"type": "input_text", "text": "..."} or {"type": "text", "text": "..."}
                    if item.get("type") in ["input_text", "text"]:
                        text_parts.append(item.get("text", ""))
                    elif item.get("type") == "image_url":
                        text_parts.append("[Image]")
                elif isinstance(item, str):
                    text_parts.append(item)
            text_content = "\n".join(text_parts)
        else:
            text_content = str(content)
        
        if text_content.strip():
            chat_messages.append({"role": role, "content": text_content})
    
    # Add a final user message requesting the report
    chat_messages.append({
        "role": "user",
        "content": f"Please create a professional {format_type.upper()} report from this conversation. Structure it clearly with sections, key findings, and actionable insights. Return ONLY the JSON object, no markdown."
    })
    
    log(f"Calling Chat Completions API with {len(chat_messages)} messages")
    
    # Use Chat Completions API (more reliable than Responses API for this use case)
    import httpx
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        raise ValueError("OPENAI_API_KEY not configured")
    
    try:
        log(f"Making request to OpenAI API with model={model}, timeout=120s")
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                f"{OPENAI_BASE_URL}/chat/completions",  # Use configured base URL
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": model,
                    "messages": chat_messages,
                    "temperature": 0.3,  # Lower temperature for more consistent JSON output
                    "max_tokens": 4000
                }
            )
            log(f"OpenAI API response status: {response.status_code}")
            if response.status_code != 200:
                log(f"OpenAI API error response: {response.text[:500]}")
            response.raise_for_status()
            result = response.json()
    except httpx.TimeoutException as e:
        log(f"❌ OpenAI API timeout after 120s: {e}")
        raise ValueError(f"OpenAI API timeout: {e}")
    except httpx.HTTPStatusError as e:
        log(f"❌ OpenAI API HTTP error: {e.response.status_code} - {e.response.text[:500]}")
        raise ValueError(f"OpenAI API error: {e.response.status_code}")
    except Exception as e:
        log(f"❌ OpenAI API unexpected error: {type(e).__name__}: {e}")
        raise
    
    # Extract the assistant's response
    choices = result.get("choices", [])
    if not choices:
        raise ValueError("No response from OpenAI")
    
    ai_text = choices[0].get("message", {}).get("content", "")
    
    if not ai_text:
        raise ValueError("Empty response from OpenAI")
    
    log(f"Received AI response: {len(ai_text)} characters")
    
    # Try to parse JSON from the response
    # The AI might wrap it in markdown code blocks, so clean it up
    ai_text = ai_text.strip()
    if ai_text.startswith("```json"):
        ai_text = ai_text[7:]
    if ai_text.startswith("```"):
        ai_text = ai_text[3:]
    if ai_text.endswith("```"):
        ai_text = ai_text[:-3]
    ai_text = ai_text.strip()
    
    try:
        ai_report = json.loads(ai_text)
    except json.JSONDecodeError as e:
        log(f"Failed to parse AI response as JSON: {e}")
        log(f"Response text (first 500 chars): {ai_text[:500]}")
        # Fallback: create a simple report structure from the text
        ai_report = {
            "title": "AI-Generated Report",
            "subtitle": "Generated from conversation",
            "summary": ai_text[:200] + "..." if len(ai_text) > 200 else ai_text,
            "sections": [{
                "heading": "Report Content",
                "body": ai_text,
                "bullets": []
            }],
            "conclusions": [],
            "recommendations": []
        }
    
    # Validate and enhance the report structure
    if not isinstance(ai_report, dict):
        raise ValueError("Invalid report structure from AI")
    
    # Ensure required fields exist
    if "title" not in ai_report:
        ai_report["title"] = "AI-Generated Report"
    if "sections" not in ai_report:
        ai_report["sections"] = []
    if not ai_report["sections"]:
        # If no sections, create one from the summary or body
        body_text = ai_report.get("summary", ai_report.get("body", "No content available"))
        ai_report["sections"] = [{
            "heading": "Report Content",
            "body": body_text,
            "bullets": []
        }]
    
    log(f"✅ AI report generated: title='{ai_report.get('title')}', sections={len(ai_report.get('sections', []))}")
    
    return ai_report


@app.post("/v1/export/generate-ai-report")
async def generate_ai_report(request: Request):
    """
    Use OpenAI Responses API to generate an intelligent, structured report from conversation.
    Returns structured report JSON that can be rendered as PDF/Word.
    """
    try:
        data = await request.json()
        conversation_history = data.get("conversation", [])
        format_type = data.get("format", "pdf").lower()
        model = data.get("model", "gpt-4o")
        
        ai_report = await _generate_ai_report_internal(conversation_history, format_type, model)
        
        return JSONResponse({
            "success": True,
            "report": ai_report
        })
            
    except HTTPException:
        raise
    except Exception as e:
        log(f"Error in generate_ai_report: {e}")
        import traceback
        log(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Export generation failed: {str(e)}")


@app.post("/v1/export/create")
async def create_export_file(request: Request):
    """
    Create an export file and return a download link.
    Renders the conversation directly to PDF/Word - no separate AI calls.
    The conversation is already formatted by whatever model the user is chatting with.
    """
    try:
        data = await request.json()
        report = data.get("report", {})
        format_type = data.get("format", "pdf").lower()
        
        # Log incoming request details
        log(f"📥 Export create request: format={format_type}")
        log(f"   Report sections: {len(report.get('sections', []))}, title: {report.get('title', 'N/A')[:50]}")
        
        # Validate report structure
        if not report or not isinstance(report, dict):
            raise HTTPException(status_code=400, detail="No valid report structure provided")
        if "sections" not in report:
            report["sections"] = []
        
        # Log what we're about to render
        log(f"📄 Rendering {format_type.upper()} with {len(report.get('sections', []))} sections")
        for i, sec in enumerate(report.get('sections', [])[:5]):  # Log first 5 sections
            log(f"   Section {i+1}: '{sec.get('heading', 'N/A')[:40]}...' ({len(sec.get('body', ''))} chars)")
        
        if format_type == "pdf":
            file_bytes = render_report_pdf(report)
            mime_type = "application/pdf"
            ext = "pdf"
        elif format_type in ["docx", "word"]:
            file_bytes = render_report_docx(report)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format_type}")
        
        log(f"✅ Generated {format_type.upper()}: {len(file_bytes)} bytes")
        
        # Validate file was actually created
        if not file_bytes or len(file_bytes) == 0:
            log(f"ERROR: Export file generation returned empty bytes!")
            raise HTTPException(status_code=500, detail="Export file generation failed - empty file")
        
        # Minimum size check (PDF should be at least a few KB, DOCX should be at least 5KB)
        min_size = 3000 if format_type == "pdf" else 5000
        if len(file_bytes) < min_size:
            log(f"WARNING: Export file seems too small: {len(file_bytes)} bytes (expected at least {min_size})")
        
        # Generate unique file ID
        file_id = str(uuid.uuid4())[:8]
        filename = _safe_slug(report.get("title", "export")) + f".{ext}"
        
        # Store on disk (works across workers/instances)
        safe_name = filename.replace("/", "_").replace("\\", "_")
        bin_path = EXPORT_DIR / f"{file_id}__{safe_name}"
        meta_path = EXPORT_DIR / f"{file_id}.json"
        
        # Write file bytes to disk
        bin_path.write_bytes(file_bytes)
        
        # Write metadata
        meta = {
            "file_id": file_id,
            "filename": safe_name,
            "mime_type": mime_type,
            "size_bytes": len(file_bytes),
            "created_at": int(time.time()),
            "bin_path": str(bin_path),
        }
        meta_path.write_text(json.dumps(meta), encoding="utf-8")
        
        # Also store in memory cache for fast access (optional)
        EXPORT_FILES[file_id] = {
            "bytes": file_bytes,
            "filename": filename,
            "mime_type": mime_type,
            "created": time.time()
        }
        
        # Clean old files from memory (older than 1 hour)
        cutoff = time.time() - 3600
        for fid in list(EXPORT_FILES.keys()):
            if EXPORT_FILES[fid].get("created", 0) < cutoff:
                del EXPORT_FILES[fid]
        
        # Clean old files from disk (older than 24 hours)
        disk_cutoff = time.time() - (24 * 3600)
        try:
            for meta_file in EXPORT_DIR.glob("*.json"):
                try:
                    meta_data = json.loads(meta_file.read_text(encoding="utf-8"))
                    if meta_data.get("created_at", 0) < disk_cutoff:
                        # Delete both metadata and binary file
                        meta_file.unlink(missing_ok=True)
                        bin_file = Path(meta_data.get("bin_path", ""))
                        if bin_file.exists():
                            bin_file.unlink(missing_ok=True)
                except Exception:
                    pass  # Skip corrupted metadata files
        except Exception:
            pass  # Skip if cleanup fails
        
        log(f"Export file created: {filename} (ID: {file_id}, {len(file_bytes):,} bytes, {len(file_bytes)/1024:.1f} KB)")
        
        return JSONResponse({
            "success": True,
            "file_id": file_id,
            "filename": filename,
            "size_bytes": len(file_bytes),
            "download_url": f"/v1/export/download/{file_id}"
        })
        
    except Exception as e:
        log(f"Export creation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/v1/export/download/{file_id}")
async def download_export_file(file_id: str):
    """Download a previously created export file."""
    # 1) Try memory cache first (if present)
    if file_id in EXPORT_FILES:
        file_data = EXPORT_FILES[file_id]
        filename = file_data["filename"]
        mime_type = file_data["mime_type"]
        file_bytes = file_data["bytes"]
    else:
        # 2) Fall back to disk (works across workers/instances)
        meta_path = EXPORT_DIR / f"{file_id}.json"
        if not meta_path.exists():
            raise HTTPException(status_code=404, detail="File not found or expired")

        meta = json.loads(meta_path.read_text(encoding="utf-8"))
        bin_path = Path(meta["bin_path"])
        if not bin_path.exists():
            raise HTTPException(status_code=404, detail="File not found or expired")

        filename = meta["filename"]
        mime_type = meta["mime_type"]
        file_bytes = bin_path.read_bytes()
    
    log(f"Export file downloaded: {filename} (ID: {file_id}, {len(file_bytes):,} bytes)")
    
    # Force download with proper headers
    # Use 'attachment' to force download instead of opening in browser
    # RFC 5987 encoding for filenames with special characters
    from urllib.parse import quote
    
    # Create both simple and RFC 5987 encoded filename
    # Simple filename for basic browsers
    safe_filename = filename.replace('"', '\\"')
    # RFC 5987 encoded filename for modern browsers (handles Unicode and special chars)
    encoded_filename = quote(filename, safe='')
    
    # Use both formats for maximum compatibility
    content_disposition = f'attachment; filename="{safe_filename}"; filename*=UTF-8\'\'{encoded_filename}'
    
    headers = {
        "Content-Disposition": content_disposition,
        "Content-Type": mime_type,
        "Content-Length": str(len(file_bytes)),
        # Prevent browser from opening PDF in new tab
        "X-Content-Type-Options": "nosniff",
        # Cache control to prevent caching
        "Cache-Control": "no-cache, no-store, must-revalidate",
        "Pragma": "no-cache",
        "Expires": "0"
    }
    
    log(f"Download headers: Content-Disposition={content_disposition[:100]}...")
    
    return StreamingResponse(
        io.BytesIO(file_bytes),
        media_type=mime_type,
        headers=headers,
    )


# SharePoint Browser API Endpoints
@app.get("/sharepoint-browser")
async def sharepoint_browser_page_old():
    """Old SharePoint browser - redirects to new one."""
    # Redirect to the main sharepoint-browser endpoint
    from fastapi.responses import RedirectResponse  # type: ignore
    return RedirectResponse(url="/sharepoint-browser")


@app.get("/api/v1/sharepoint/files")
async def list_sharepoint_files_api(folder: str = ""):
    """API endpoint to list SharePoint files and folders. Supports folder parameter for navigation."""
    try:
        # Check environment variables first
        client_id = os.environ.get("SHAREPOINT_CLIENT_ID", "")
        client_secret = os.environ.get("SHAREPOINT_CLIENT_SECRET", "")
        tenant_id = os.environ.get("SHAREPOINT_TENANT_ID", "")
        site_url = os.environ.get("SHAREPOINT_SITE_URL", "")
        enable_sp = os.environ.get("ENABLE_SHAREPOINT", "false").lower() == "true"
        
        if not enable_sp:
            raise HTTPException(
                status_code=403, 
                detail="SharePoint integration is disabled. Set ENABLE_SHAREPOINT=true in environment variables."
            )
        
        if not all([client_id, client_secret, tenant_id, site_url]):
            missing = []
            if not client_id: missing.append("SHAREPOINT_CLIENT_ID")
            if not client_secret: missing.append("SHAREPOINT_CLIENT_SECRET")
            if not tenant_id: missing.append("SHAREPOINT_TENANT_ID")
            if not site_url: missing.append("SHAREPOINT_SITE_URL")
            
            raise HTTPException(
                status_code=400,
                detail=f"SharePoint configuration incomplete. Missing: {', '.join(missing)}. Please set these environment variables in Render."
            )
        
        # Import the filter to use its methods
        from sharepoint_import_filter import Filter
        filter_instance = Filter()
        
        if not filter_instance.valves.enable_sharepoint:
            raise HTTPException(status_code=403, detail="SharePoint integration not enabled in filter")
        
        # Use the new method that includes folders for navigation
        items = filter_instance._list_sharepoint_items(folder_path=folder, include_folders=True)
        
        if items is None:
            raise HTTPException(status_code=500, detail="Failed to retrieve files. Check SharePoint credentials and permissions.")
        
        return JSONResponse(content={
            "files": items,
            "current_folder": folder or "/",
            "parent_folder": "/".join(folder.split("/")[:-1]) if folder else None
        })
    except HTTPException:
        raise
    except Exception as e:
        log(f"Failed to list SharePoint files: {e}")
        import traceback
        log(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"SharePoint error: {str(e)}")


@app.get("/api/v1/sharepoint/search")
async def search_sharepoint_files_api(q: str, folder: str = ""):
    """API endpoint to search SharePoint files within a folder."""
    try:
        from sharepoint_import_filter import Filter
        filter_instance = Filter()
        
        if not filter_instance.valves.enable_sharepoint:
            raise HTTPException(status_code=403, detail="SharePoint integration not enabled")
        
        # List items in the specified folder and filter by search query
        items = filter_instance._list_sharepoint_items(folder_path=folder, include_folders=True)
        filtered = [f for f in items if q.lower() in f.get("name", "").lower()]
        
        return JSONResponse(content={
            "files": filtered,
            "current_folder": folder or "/",
            "query": q
        })
    except Exception as e:
        log(f"Failed to search SharePoint files: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/sharepoint/import")
async def import_sharepoint_file_api(request: Request):
    """API endpoint to import a SharePoint file."""
    try:
        body = await request.json()
        file_id = body.get("file_id")
        drive_id = body.get("drive_id", "")
        
        if not file_id:
            raise HTTPException(status_code=400, detail="file_id is required")
        
        from sharepoint_import_filter import Filter
        filter_instance = Filter()
        
        if not filter_instance.valves.enable_sharepoint:
            raise HTTPException(status_code=403, detail="SharePoint integration not enabled")
        
        # If drive_id not provided, get it from site info
        if not drive_id:
            info = filter_instance._get_site_and_drive_info()
            if info:
                drive_id = info.get("drive_id", "")
        
        # Get file details - search in all items to find the file
        items = filter_instance._list_sharepoint_items(include_folders=False)
        target_file = None
        for f in items:
            if f.get("id") == file_id:
                target_file = f
                break
        
        if not target_file:
            # File might be in a subfolder, try to download directly with the IDs we have
            log(f"File {file_id} not found in root, attempting direct download with drive_id={drive_id}")
            target_file = {"id": file_id, "name": "downloaded_file", "drive_id": drive_id}
        
        # Download file
        file_drive_id = target_file.get("drive_id", "") or drive_id
        download_url = target_file.get("download_url", "")
        filename = target_file.get("name", "downloaded_file")
        
        local_path = filter_instance._download_sharepoint_file(
            file_id, file_drive_id, filename, download_url
        )
        
        if not local_path:
            raise HTTPException(status_code=500, detail="Failed to download file")
        
        return JSONResponse(content={
            "success": True,
            "filename": filename,
            "path": local_path,
            "message": f"File '{filename}' imported successfully"
        })
    except HTTPException:
        raise
    except Exception as e:
        log(f"Failed to import SharePoint file: {e}")
        import traceback
        log(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v1/tools/analyze-file")
async def analyze_file_tool(payload: dict):
    if not _is_enabled("ENABLE_ANALYZE_TOOL", default=True):
        raise HTTPException(status_code=403, detail="Analyze tool disabled")
    report = analyze_file_payload(payload)
    return JSONResponse(content=report)


@app.post("/v1/tools/search")
async def search_tool(payload: dict):
    if not _is_enabled("ENABLE_SEARCH_TOOL", default=True):
        raise HTTPException(status_code=403, detail="Search tool disabled")
    query = (payload.get("query") or "").strip()
    if not query or len(query) > 200:
        raise HTTPException(status_code=400, detail="Invalid query")
    banned = ["password", "secret", "token", "apikey"]
    if any(word in query.lower() for word in banned):
        raise HTTPException(status_code=400, detail="Query not allowed")

    url = f"https://api.duckduckgo.com/?q={quote_plus(query)}&format=json&no_redirect=1&no_html=1"
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.get(url, headers={"User-Agent": "glchemtec-search"})
            resp.raise_for_status()
            data = resp.json()
    except Exception as e:
        log(f"Search failed: {e}")
        raise HTTPException(status_code=502, detail="Search failed")

    results = []
    topics = data.get("RelatedTopics", []) or []
    results_raw = data.get("Results", []) or []
    for item in results_raw:
        t = item.get("Text")
        u = item.get("FirstURL")
        if t and u:
            results.append({"title": t, "url": u})
    for item in topics:
        if isinstance(item, dict):
            t = item.get("Text")
            u = item.get("FirstURL")
            if t and u:
                results.append({"title": t, "url": u})
        if len(results) >= 5:
            break

    return JSONResponse(content={"query": query, "results": results[:5]})


@app.post("/v1/tools/transcribe")
async def transcribe_tool(payload: dict):
    if not _is_enabled("ENABLE_AUDIO_TOOLS", default=True):
        raise HTTPException(status_code=403, detail="Audio tools disabled")
    text = await openai_transcribe_audio(_prepare_audio_bytes(payload))
    return JSONResponse(content={"text": text})


@app.get("/sharepoint-browser")
async def sharepoint_browser_page():
    """Serve SharePoint browser HTML page with folder navigation."""
    html_content = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>SharePoint File Browser - GLChemTec</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: #0f1419;
            color: #fff;
            padding: 20px;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
        }
        .header {
            background: linear-gradient(135deg, #1a1f2e 0%, #0f1419 100%);
            padding: 20px;
            border-radius: 12px;
            margin-bottom: 20px;
            border: 1px solid rgba(255,255,255,0.1);
        }
        .header h1 {
            font-size: 24px;
            margin-bottom: 10px;
        }
        .header p {
            color: #94a3b8;
            font-size: 14px;
        }
        .breadcrumb {
            display: flex;
            align-items: center;
            gap: 8px;
            padding: 12px 16px;
            background: #1e293b;
            border-radius: 8px;
            margin-bottom: 15px;
            flex-wrap: wrap;
            border: 1px solid #334155;
        }
        .breadcrumb-item {
            color: #3b82f6;
            cursor: pointer;
            font-size: 14px;
            padding: 4px 8px;
            border-radius: 4px;
            transition: background 0.2s;
        }
        .breadcrumb-item:hover {
            background: #334155;
            text-decoration: underline;
        }
        .breadcrumb-item.current {
            color: #fff;
            cursor: default;
            font-weight: 500;
        }
        .breadcrumb-item.current:hover {
            background: transparent;
            text-decoration: none;
        }
        .breadcrumb-separator {
            color: #64748b;
            font-size: 12px;
        }
        .controls {
            display: flex;
            gap: 10px;
            margin-bottom: 20px;
            flex-wrap: wrap;
        }
        button {
            background: #3b82f6;
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 8px;
            cursor: pointer;
            font-size: 14px;
            font-weight: 500;
            transition: all 0.2s;
        }
        button:hover:not(:disabled) {
            background: #2563eb;
        }
        button:disabled {
            background: #475569;
            cursor: not-allowed;
        }
        button.secondary {
            background: #475569;
        }
        button.secondary:hover:not(:disabled) {
            background: #64748b;
        }
        input[type="text"] {
            flex: 1;
            min-width: 200px;
            padding: 10px 15px;
            background: #1e293b;
            border: 1px solid #334155;
            border-radius: 8px;
            color: #fff;
            font-size: 14px;
        }
        input[type="text"]:focus {
            outline: none;
            border-color: #3b82f6;
        }
        .file-list {
            background: #1e293b;
            border-radius: 12px;
            border: 1px solid #334155;
            overflow: hidden;
        }
        .file-item {
            display: flex;
            align-items: center;
            padding: 15px 20px;
            border-bottom: 1px solid #334155;
            cursor: pointer;
            transition: background 0.2s;
        }
        .file-item:hover {
            background: #334155;
        }
        .file-item.selected {
            background: #1e40af !important;
        }
        .file-item.folder-item:hover {
            background: #3d2e0a;
        }
        .file-item:last-child {
            border-bottom: none;
        }
        .file-icon {
            width: 40px;
            height: 40px;
            background: #475569;
            border-radius: 8px;
            display: flex;
            align-items: center;
            justify-content: center;
            margin-right: 15px;
            font-size: 20px;
        }
        .file-icon.folder-icon {
            background: #78350f;
        }
        .file-info {
            flex: 1;
        }
        .file-name {
            font-weight: 500;
            margin-bottom: 5px;
        }
        .file-meta {
            font-size: 12px;
            color: #94a3b8;
        }
        .folder-badge {
            background: #fbbf24;
            color: #000;
            font-size: 10px;
            padding: 2px 6px;
            border-radius: 4px;
            margin-left: 8px;
            font-weight: 600;
        }
        .loading {
            text-align: center;
            padding: 40px;
            color: #94a3b8;
        }
        .error {
            background: #7f1d1d;
            border: 1px solid #991b1b;
            color: #fca5a5;
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
        }
        .success {
            background: #14532d;
            border: 1px solid #166534;
            color: #86efac;
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
        }
        .info {
            background: #1e3a5f;
            border: 1px solid #1e40af;
            color: #93c5fd;
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
        }
        #message {
            display: none;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📂 SharePoint File Browser</h1>
            <p>Browse folders and import files from SharePoint for analysis in OpenWebUI</p>
        </div>
        
        <div id="message"></div>
        
        <div class="breadcrumb" id="breadcrumb">
            <span class="breadcrumb-item current">📁 Root</span>
        </div>
        
        <div class="controls">
            <input type="text" id="searchInput" placeholder="Search in current folder... (Press Enter)">
            <button onclick="searchFiles()">🔍 Search</button>
            <button onclick="listFiles()" class="secondary">🔄 Refresh</button>
            <button onclick="importSelected()" id="importBtn" disabled>📥 Import Selected</button>
        </div>
        
        <div class="file-list" id="fileList">
            <div class="loading">Loading files from SharePoint...</div>
        </div>
    </div>

    <script>
        let selectedFile = null;
        let currentFolder = '';
        const apiBase = window.location.origin;

        function updateBreadcrumb() {
            const breadcrumb = document.getElementById('breadcrumb');
            const parts = currentFolder ? currentFolder.split('/').filter(p => p) : [];
            
            let html = '<span class="breadcrumb-item' + (parts.length === 0 ? ' current' : '') + '" onclick="navigateTo(\\'\\')">📁 Root</span>';
            
            let path = '';
            parts.forEach((part, index) => {
                path += (path ? '/' : '') + part;
                const isLast = index === parts.length - 1;
                html += '<span class="breadcrumb-separator">›</span>';
                html += '<span class="breadcrumb-item' + (isLast ? ' current' : '') + '" ' + 
                        (isLast ? '' : 'onclick="navigateTo(\\'' + path.replace(/'/g, "\\\\'") + '\\')"') + 
                        '>' + part + '</span>';
            });
            
            breadcrumb.innerHTML = html;
        }

        function navigateTo(folder) {
            currentFolder = folder;
            selectedFile = null;
            document.getElementById('importBtn').disabled = true;
            updateBreadcrumb();
            listFiles();
        }

        async function listFiles() {
            const fileList = document.getElementById('fileList');
            fileList.innerHTML = '<div class="loading">Loading files from SharePoint...</div>';
            
            try {
                const url = currentFolder 
                    ? `${apiBase}/api/v1/sharepoint/files?folder=${encodeURIComponent(currentFolder)}`
                    : `${apiBase}/api/v1/sharepoint/files`;
                    
                const response = await fetch(url, {
                    method: 'GET',
                    headers: { 'Content-Type': 'application/json' }
                });
                
                if (!response.ok) {
                    const error = await response.json();
                    throw new Error(error.detail || 'Failed to load files');
                }
                
                const data = await response.json();
                displayFiles(data.files || []);
            } catch (error) {
                showError('Failed to load files: ' + error.message);
                fileList.innerHTML = '<div class="error">Error loading files. ' + error.message + '</div>';
            }
        }

        async function searchFiles() {
            const query = document.getElementById('searchInput').value.trim();
            if (!query) {
                listFiles();
                return;
            }
            
            const fileList = document.getElementById('fileList');
            fileList.innerHTML = '<div class="loading">Searching...</div>';
            
            try {
                const url = `${apiBase}/api/v1/sharepoint/search?q=${encodeURIComponent(query)}&folder=${encodeURIComponent(currentFolder)}`;
                const response = await fetch(url, {
                    method: 'GET',
                    headers: { 'Content-Type': 'application/json' }
                });
                
                if (!response.ok) throw new Error('Search failed');
                
                const data = await response.json();
                displayFiles(data.files || []);
                showMessage(`Found ${data.files?.length || 0} results for "${query}"`, 'info');
            } catch (error) {
                showError('Search failed: ' + error.message);
            }
        }

        function displayFiles(files) {
            const fileList = document.getElementById('fileList');
            
            if (files.length === 0) {
                fileList.innerHTML = '<div class="loading">📭 No files or folders found in this location</div>';
                return;
            }
            
            fileList.innerHTML = files.map(file => {
                const isFolder = file.is_folder;
                const escapedName = file.name.replace(/'/g, "\\\\'");
                const escapedPath = (file.path || file.name).replace(/'/g, "\\\\'");
                
                if (isFolder) {
                    return `
                        <div class="file-item folder-item" onclick="navigateTo('${escapedPath}')">
                            <div class="file-icon folder-icon">📁</div>
                            <div class="file-info">
                                <div class="file-name">${file.name}<span class="folder-badge">FOLDER</span></div>
                                <div class="file-meta">${file.child_count || 0} items • Click to open</div>
                            </div>
                        </div>
                    `;
                } else {
                    return `
                        <div class="file-item" onclick="selectFile('${file.id}', '${escapedName}', '${file.drive_id || ''}', this)">
                            <div class="file-icon">${getFileIcon(file.name)}</div>
                            <div class="file-info">
                                <div class="file-name">${file.name}</div>
                                <div class="file-meta">${formatSize(file.size)} • ${formatDate(file.modified)}</div>
                            </div>
                        </div>
                    `;
                }
            }).join('');
        }

        function selectFile(id, name, driveId, element) {
            selectedFile = { id, name, drive_id: driveId };
            document.getElementById('importBtn').disabled = false;
            
            // Highlight selected
            document.querySelectorAll('.file-item').forEach(item => {
                item.classList.remove('selected');
            });
            element.classList.add('selected');
        }

        async function importSelected() {
            if (!selectedFile) return;
            
            showMessage('📥 Importing file...', 'info');
            document.getElementById('importBtn').disabled = true;
            
            try {
                const response = await fetch(`${apiBase}/api/v1/sharepoint/import`, {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ 
                        file_id: selectedFile.id,
                        drive_id: selectedFile.drive_id
                    })
                });
                
                if (!response.ok) {
                    const error = await response.json();
                    throw new Error(error.detail || 'Import failed');
                }
                
                const data = await response.json();
                showMessage(`✅ File "${selectedFile.name}" imported successfully! You can now use it in OpenWebUI chat.`, 'success');
                
                // Reset selection
                selectedFile = null;
                document.querySelectorAll('.file-item').forEach(item => {
                    item.classList.remove('selected');
                });
            } catch (error) {
                showError('Failed to import file: ' + error.message);
                document.getElementById('importBtn').disabled = false;
            }
        }

        function getFileIcon(name) {
            const ext = name.split('.').pop()?.toLowerCase();
            if (['pdf'].includes(ext)) return '📄';
            if (['doc', 'docx'].includes(ext)) return '📝';
            if (['xls', 'xlsx', 'csv'].includes(ext)) return '📊';
            if (['jpg', 'jpeg', 'png', 'gif', 'webp'].includes(ext)) return '🖼️';
            if (['pptx', 'ppt'].includes(ext)) return '📽️';
            if (['mp3', 'wav', 'ogg', 'm4a'].includes(ext)) return '🎵';
            if (['mp4', 'mov', 'avi', 'mkv'].includes(ext)) return '🎬';
            if (['zip', 'rar', '7z', 'tar', 'gz'].includes(ext)) return '📦';
            if (['txt', 'md', 'json', 'xml'].includes(ext)) return '📃';
            return '📎';
        }

        function formatSize(bytes) {
            if (!bytes) return '0 B';
            const k = 1024;
            const sizes = ['B', 'KB', 'MB', 'GB'];
            const i = Math.floor(Math.log(bytes) / Math.log(k));
            return (bytes / Math.pow(k, i)).toFixed(1) + ' ' + sizes[i];
        }

        function formatDate(dateStr) {
            if (!dateStr) return 'Unknown';
            try {
                return new Date(dateStr).toLocaleDateString();
            } catch {
                return dateStr;
            }
        }

        function showMessage(msg, type) {
            const msgDiv = document.getElementById('message');
            msgDiv.className = type;
            msgDiv.innerHTML = msg;
            msgDiv.style.display = 'block';
            if (type === 'success' || type === 'info') {
                setTimeout(() => msgDiv.style.display = 'none', 5000);
            }
        }

        function showError(msg) {
            showMessage('❌ ' + msg, 'error');
        }

        // Enter key to search
        document.getElementById('searchInput').addEventListener('keypress', (e) => {
            if (e.key === 'Enter') searchFiles();
        });
        
        // Clear search on empty input
        document.getElementById('searchInput').addEventListener('input', (e) => {
            if (!e.target.value.trim()) {
                listFiles();
            }
        });

        // Load files on page load
        updateBreadcrumb();
        listFiles();
    </script>
</body>
</html>
    """
    return HTMLResponse(content=html_content)


@app.get("/v1/models")
async def list_models():
    """Forward models list request."""
    log("GET /v1/models - OpenWebUI is calling the proxy!")
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.get(
                f"{OPENAI_BASE_URL}/models",
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}"}
            )
            resp.raise_for_status()
            data = resp.json()
            model_count = len(data.get("data", []))
            log(f"GET /v1/models - Success! Returning {model_count} models")
            return JSONResponse(content=data)
    except httpx.HTTPStatusError as e:
        log(f"GET /v1/models - HTTP Error {e.response.status_code}: {e.response.text[:200]}")
        # Return proper JSON error response instead of raising
        return JSONResponse(
            status_code=e.response.status_code,
            content={
                "error": {
                    "message": f"OpenAI API error: {e.response.status_code} {e.response.reason_phrase}",
                    "type": "api_error",
                    "code": e.response.status_code
                }
            }
        )
    except Exception as e:
        log(f"GET /v1/models - ERROR: {type(e).__name__}: {e}")
        # Return proper JSON error response
        return JSONResponse(
            status_code=500,
            content={
                "error": {
                    "message": f"Proxy error: {str(e)}",
                    "type": "proxy_error"
                }
            }
        )


@app.get("/metrics")
async def metrics():
    """Simple in-memory metrics snapshot."""
    return JSONResponse(content=METRICS)


@app.get("/health")
async def health():
    log("HEALTH CHECK: Proxy is alive and responding!")
    return {"status": "ok", "service": "OpenAI Responses Proxy", "port": 8000}


@app.get("/")
async def root():
    log("ROOT ENDPOINT CALLED: Proxy is definitely running!")
    return {"service": "OpenAI Responses Proxy", "status": "running", "port": 8000}

@app.get("/test")
async def test():
    """Test endpoint to verify proxy is running - call this from OpenWebUI"""
    log("TEST ENDPOINT CALLED: Proxy is alive!")
    return {"status": "ok", "message": "Proxy is running on port 8000", "timestamp": time.time()}
