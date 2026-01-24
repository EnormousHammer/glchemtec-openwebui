"""
title: Document Filter
author: GLChemTec
version: 1.1
description: Extracts text and images from DOCX, XLSX, CSV, ChemDraw, and other document formats for OpenAI Vision analysis.
"""

import os
import base64
import tempfile
import zipfile
import io
import csv
import json
from typing import Optional, List, Dict, Any

from pydantic import BaseModel, Field

# Optional imports
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

try:
    import xlrd
    XLS_AVAILABLE = True
except ImportError:
    XLS_AVAILABLE = False


class Filter:
    class Valves(BaseModel):
        priority: int = Field(default=1, description="Filter priority (runs after PPT filter)")
        enabled: bool = Field(default=True, description="Enable document processing")
        debug: bool = Field(default=True, description="Enable debug logging")
        
        # Limits
        max_text_chars: int = Field(default=100000, description="Max characters to extract from text")
        max_images: int = Field(default=50, description="Max images to extract per document")
        max_image_size_mb: float = Field(default=10.0, description="Max size per image (MB)")

    def __init__(self):
        self.valves = self.Valves()

    def _log(self, msg: str) -> None:
        if self.valves.debug:
            print(f"[DOCUMENT-FILTER] {msg}")

    # -------------------------
    # File discovery (same as ppt_pdf_filter)
    # -------------------------
    def _get_file_path(self, file_obj: Dict[str, Any]) -> str:
        if not isinstance(file_obj, dict):
            return ""
        if isinstance(file_obj.get("file"), dict):
            f = file_obj["file"]
            path = (f.get("path") or "").strip()
            if path:
                return path
            if isinstance(f.get("meta"), dict):
                return (f["meta"].get("path") or "").strip()
        return (file_obj.get("path") or "").strip()

    def _get_file_name(self, file_obj: Dict[str, Any]) -> str:
        if not isinstance(file_obj, dict):
            return ""
        if isinstance(file_obj.get("file"), dict):
            f = file_obj["file"]
            name = f.get("filename") or f.get("name") or ""
            if not name and isinstance(f.get("meta"), dict):
                name = f["meta"].get("name") or ""
            return (name or "").lower().strip()
        return ((file_obj.get("name") or file_obj.get("filename") or "")).lower().strip()

    def _extract_all_files(self, body: dict, messages: list) -> List[Dict[str, Any]]:
        all_files: List[Dict[str, Any]] = []

        if isinstance(body.get("files"), list):
            all_files.extend(body["files"])

        for msg in messages:
            if isinstance(msg.get("files"), list):
                all_files.extend(msg["files"])
            if isinstance(msg.get("attachments"), list):
                all_files.extend(msg["attachments"])
            if isinstance(msg.get("sources"), list):
                for source_obj in msg["sources"]:
                    if isinstance(source_obj, dict):
                        source = source_obj.get("source", {})
                        if source.get("type") == "file" and isinstance(source.get("file"), dict):
                            all_files.append({"file": source["file"]})

        # Dedup by path
        seen = set()
        unique = []
        for f in all_files:
            p = self._get_file_path(f)
            if p and p not in seen:
                seen.add(p)
                unique.append(f)
        return unique

    def _extract_text_content(self, content: Any) -> str:
        if isinstance(content, str):
            return content
        if isinstance(content, list):
            parts = []
            for item in content:
                if isinstance(item, dict) and item.get("type") == "text":
                    parts.append(item.get("text", ""))
            return "\n".join([p for p in parts if p]).strip()
        return str(content) if content else ""

    # -------------------------
    # Image utilities
    # -------------------------
    def _image_to_data_url(self, image_data: bytes, ext: str) -> Optional[str]:
        """Convert image bytes to data URL."""
        try:
            max_bytes = int(self.valves.max_image_size_mb * 1024 * 1024)
            if len(image_data) > max_bytes:
                self._log(f"Image too large ({len(image_data)} bytes), skipping")
                return None
            
            ext = ext.lower().strip('.')
            if ext in ('png',):
                mime = 'image/png'
            elif ext in ('jpg', 'jpeg'):
                mime = 'image/jpeg'
            elif ext in ('gif',):
                mime = 'image/gif'
            elif ext in ('webp',):
                mime = 'image/webp'
            elif ext in ('bmp',):
                mime = 'image/bmp'
            else:
                mime = 'image/png'  # Default
            
            b64 = base64.b64encode(image_data).decode('utf-8')
            return f"data:{mime};base64,{b64}"
        except Exception as e:
            self._log(f"Error converting image: {e}")
            return None

    # -------------------------
    # DOCX extraction
    # -------------------------
    def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text and images from DOCX."""
        result = {"text": "", "images": []}
        
        if not DOCX_AVAILABLE:
            self._log("python-docx not available, extracting text only via zipfile")
        
        try:
            # Extract images from zip structure
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Get all image files
                image_files = [f for f in docx_zip.namelist() 
                              if f.startswith('word/media/') and 
                              any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'])]
                
                for img_file in image_files[:self.valves.max_images]:
                    try:
                        img_data = docx_zip.read(img_file)
                        ext = os.path.splitext(img_file)[1]
                        data_url = self._image_to_data_url(img_data, ext)
                        if data_url:
                            result["images"].append({
                                "type": "image_url",
                                "image_url": {"url": data_url}
                            })
                    except Exception as e:
                        self._log(f"Error extracting image {img_file}: {e}")
                
                self._log(f"Extracted {len(result['images'])} images from DOCX")
            
            # Extract text
            if DOCX_AVAILABLE:
                doc = Document(file_path)
                paragraphs = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())
                
                # Extract tables
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells]
                        table_text.append(" | ".join(row_cells))
                    if table_text:
                        paragraphs.append("\n[Table]\n" + "\n".join(table_text))
                
                result["text"] = "\n\n".join(paragraphs)[:self.valves.max_text_chars]
            else:
                # Fallback: extract text from XML
                with zipfile.ZipFile(file_path, 'r') as docx_zip:
                    if 'word/document.xml' in docx_zip.namelist():
                        xml_content = docx_zip.read('word/document.xml').decode('utf-8')
                        import re
                        # Simple text extraction from XML
                        text_parts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                        result["text"] = ' '.join(text_parts)[:self.valves.max_text_chars]
            
            self._log(f"Extracted {len(result['text'])} chars of text from DOCX")
            
        except Exception as e:
            self._log(f"Error extracting DOCX: {e}")
        
        return result

    # -------------------------
    # XLSX extraction
    # -------------------------
    def _extract_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Extract text and images from XLSX."""
        result = {"text": "", "images": []}
        
        try:
            # Extract images from zip structure
            with zipfile.ZipFile(file_path, 'r') as xlsx_zip:
                image_files = [f for f in xlsx_zip.namelist() 
                              if f.startswith('xl/media/') and 
                              any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'])]
                
                for img_file in image_files[:self.valves.max_images]:
                    try:
                        img_data = xlsx_zip.read(img_file)
                        ext = os.path.splitext(img_file)[1]
                        data_url = self._image_to_data_url(img_data, ext)
                        if data_url:
                            result["images"].append({
                                "type": "image_url",
                                "image_url": {"url": data_url}
                            })
                    except Exception as e:
                        self._log(f"Error extracting image {img_file}: {e}")
                
                self._log(f"Extracted {len(result['images'])} images from XLSX")
            
            # Extract cell data
            if XLSX_AVAILABLE:
                wb = load_workbook(file_path, data_only=True)
                sheets_text = []
                
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    sheet_data = [f"=== Sheet: {sheet_name} ==="]
                    
                    for row in sheet.iter_rows(values_only=True):
                        if any(cell is not None for cell in row):
                            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                            sheet_data.append(row_text)
                    
                    if len(sheet_data) > 1:  # More than just header
                        sheets_text.append("\n".join(sheet_data))
                
                result["text"] = "\n\n".join(sheets_text)[:self.valves.max_text_chars]
            
            self._log(f"Extracted {len(result['text'])} chars of text from XLSX")
            
        except Exception as e:
            self._log(f"Error extracting XLSX: {e}")
        
        return result

    # -------------------------
    # XLS (legacy) extraction
    # -------------------------
    def _extract_xls(self, file_path: str) -> Dict[str, Any]:
        """Extract text from legacy XLS."""
        result = {"text": "", "images": []}
        
        if not XLS_AVAILABLE:
            self._log("xlrd not available, cannot extract XLS")
            return result
        
        try:
            wb = xlrd.open_workbook(file_path)
            sheets_text = []
            
            for sheet_idx in range(wb.nsheets):
                sheet = wb.sheet_by_index(sheet_idx)
                sheet_data = [f"=== Sheet: {sheet.name} ==="]
                
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    if any(cell for cell in row):
                        row_text = " | ".join(str(cell) for cell in row)
                        sheet_data.append(row_text)
                
                if len(sheet_data) > 1:
                    sheets_text.append("\n".join(sheet_data))
            
            result["text"] = "\n\n".join(sheets_text)[:self.valves.max_text_chars]
            self._log(f"Extracted {len(result['text'])} chars from XLS")
            
        except Exception as e:
            self._log(f"Error extracting XLS: {e}")
        
        return result

    # -------------------------
    # CSV/TSV extraction
    # -------------------------
    def _extract_csv(self, file_path: str, delimiter: str = ',') -> Dict[str, Any]:
        """Extract text from CSV/TSV."""
        result = {"text": "", "images": []}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f, delimiter=delimiter)
                rows = []
                for row in reader:
                    if any(cell.strip() for cell in row):
                        rows.append(" | ".join(row))
                
                result["text"] = "\n".join(rows)[:self.valves.max_text_chars]
            
            self._log(f"Extracted {len(result['text'])} chars from CSV/TSV")
            
        except Exception as e:
            self._log(f"Error extracting CSV: {e}")
        
        return result

    # -------------------------
    # Text file extraction
    # -------------------------
    def _extract_text_file(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT/MD/JSON files."""
        result = {"text": "", "images": []}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                result["text"] = f.read()[:self.valves.max_text_chars]
            
            self._log(f"Extracted {len(result['text'])} chars from text file")
            
        except Exception as e:
            self._log(f"Error extracting text file: {e}")
        
        return result

    # -------------------------
    # ChemDraw CDX/CDXML extraction
    # -------------------------
    def _extract_chemdraw(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Extract structure data and embedded images from ChemDraw files."""
        result = {"text": "", "images": []}
        
        try:
            is_cdxml = file_name.endswith('.cdxml')
            
            if is_cdxml:
                # CDXML is XML-based, can extract text directly
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                
                import re
                
                # Extract molecule names
                names = re.findall(r'<s[^>]*>([^<]+)</s>', content)
                
                # Extract text annotations
                texts = re.findall(r'<t[^>]*>([^<]*(?:<s[^>]*>[^<]*</s>[^<]*)*)</t>', content)
                
                # Extract chemical formulas if present
                formulas = re.findall(r'Formula="([^"]+)"', content)
                
                # Extract molecular weights
                mol_weights = re.findall(r'MolecularWeight="([^"]+)"', content)
                
                # Build structured output
                parts = ["=== ChemDraw Structure Data ===\n"]
                
                if formulas:
                    parts.append("**Molecular Formulas:**")
                    for f in set(formulas):
                        parts.append(f"  - {f}")
                    parts.append("")
                
                if mol_weights:
                    parts.append("**Molecular Weights:**")
                    for mw in set(mol_weights):
                        parts.append(f"  - {mw}")
                    parts.append("")
                
                # Clean and add text content
                all_text = []
                for t in texts:
                    # Remove XML tags from text
                    clean = re.sub(r'<[^>]+>', '', t).strip()
                    if clean and len(clean) > 1:
                        all_text.append(clean)
                
                if all_text:
                    parts.append("**Text/Labels:**")
                    for t in set(all_text):
                        parts.append(f"  - {t}")
                    parts.append("")
                
                # Extract embedded PNG images from CDXML
                png_matches = re.findall(r'<embedded[^>]*type="image/png"[^>]*>([^<]+)</embedded>', content)
                for i, png_b64 in enumerate(png_matches[:self.valves.max_images]):
                    try:
                        # Clean base64 string
                        clean_b64 = png_b64.strip().replace('\n', '').replace('\r', '').replace(' ', '')
                        data_url = f"data:image/png;base64,{clean_b64}"
                        result["images"].append({
                            "type": "image_url",
                            "image_url": {"url": data_url}
                        })
                    except Exception as e:
                        self._log(f"Error extracting embedded image {i}: {e}")
                
                result["text"] = "\n".join(parts)[:self.valves.max_text_chars]
                self._log(f"Extracted {len(result['text'])} chars and {len(result['images'])} images from CDXML")
                
            else:
                # CDX is binary format - extract what we can
                with open(file_path, 'rb') as f:
                    content = f.read()
                
                # Try to find text strings in binary
                import re
                # Look for ASCII text sequences (chemical names, labels)
                text_matches = re.findall(rb'[\x20-\x7E]{4,100}', content)
                
                # Filter for chemistry-relevant strings
                chem_keywords = ['mol', 'atom', 'bond', 'ring', 'chain', 'formula', 'weight', 
                                'reagent', 'product', 'yield', 'temp', 'time', 'equiv',
                                'mmol', 'mg', 'ml', 'hr', 'min', 'rt', 'reflux']
                
                relevant_texts = []
                for t in text_matches:
                    try:
                        decoded = t.decode('ascii', errors='ignore').strip()
                        # Keep if it looks like chemistry content
                        if any(kw in decoded.lower() for kw in chem_keywords) or \
                           re.match(r'^[A-Z][a-z]?\d*', decoded) or \
                           len(decoded) > 10:
                            relevant_texts.append(decoded)
                    except:
                        pass
                
                # Look for embedded PNG in binary CDX
                png_header = b'\x89PNG\r\n\x1a\n'
                png_end = b'IEND\xaeB`\x82'
                
                pos = 0
                while pos < len(content) and len(result["images"]) < self.valves.max_images:
                    start = content.find(png_header, pos)
                    if start == -1:
                        break
                    end = content.find(png_end, start)
                    if end == -1:
                        break
                    end += len(png_end)
                    
                    png_data = content[start:end]
                    if len(png_data) < self.valves.max_image_size_mb * 1024 * 1024:
                        b64 = base64.b64encode(png_data).decode('utf-8')
                        result["images"].append({
                            "type": "image_url",
                            "image_url": {"url": f"data:image/png;base64,{b64}"}
                        })
                    pos = end
                
                parts = ["=== ChemDraw Binary File ===\n"]
                parts.append("(Binary CDX format - limited text extraction)\n")
                
                if relevant_texts:
                    parts.append("**Extracted Text/Labels:**")
                    for t in list(set(relevant_texts))[:50]:  # Limit to 50 unique strings
                        parts.append(f"  - {t}")
                
                result["text"] = "\n".join(parts)[:self.valves.max_text_chars]
                self._log(f"Extracted {len(result['text'])} chars and {len(result['images'])} images from CDX")
        
        except Exception as e:
            self._log(f"Error extracting ChemDraw file: {e}")
            import traceback
            self._log(traceback.format_exc())
        
        return result

    # -------------------------
    # MOL/SDF structure file extraction
    # -------------------------
    def _extract_mol_sdf(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Extract structure data from MOL/SDF files."""
        result = {"text": "", "images": []}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            parts = [f"=== Chemical Structure File: {file_name} ===\n"]
            
            if file_name.endswith('.sdf'):
                # SDF can contain multiple molecules separated by $$$$
                molecules = content.split('$$$$')
                parts.append(f"**Contains {len([m for m in molecules if m.strip()])} molecule(s)**\n")
                
                for i, mol in enumerate(molecules[:10], 1):  # Limit to first 10
                    if not mol.strip():
                        continue
                    
                    lines = mol.strip().split('\n')
                    if lines:
                        # First line is usually molecule name
                        mol_name = lines[0].strip()
                        if mol_name:
                            parts.append(f"**Molecule {i}:** {mol_name}")
                        
                        # Look for data fields (lines starting with > <FIELD_NAME>)
                        import re
                        fields = re.findall(r'>\s*<([^>]+)>\s*\n([^\n>]+)', mol)
                        for field_name, field_value in fields:
                            parts.append(f"  {field_name}: {field_value.strip()}")
                        parts.append("")
            else:
                # Single MOL file
                lines = content.split('\n')
                if lines:
                    mol_name = lines[0].strip()
                    if mol_name:
                        parts.append(f"**Molecule Name:** {mol_name}")
                    
                    # Parse counts line (line 4 in V2000 format)
                    if len(lines) > 3:
                        counts_line = lines[3].strip()
                        # Format: aaabbblllfffcccsssxxxrrrpppiiimmmvvvvvv
                        # aaa = number of atoms, bbb = number of bonds
                        try:
                            num_atoms = int(counts_line[0:3].strip())
                            num_bonds = int(counts_line[3:6].strip())
                            parts.append(f"**Atoms:** {num_atoms}")
                            parts.append(f"**Bonds:** {num_bonds}")
                        except:
                            pass
            
            # Include raw content for AI analysis
            parts.append("\n**Raw Structure Data:**")
            parts.append("```")
            parts.append(content[:5000])  # Limit raw content
            if len(content) > 5000:
                parts.append("... (truncated)")
            parts.append("```")
            
            result["text"] = "\n".join(parts)[:self.valves.max_text_chars]
            self._log(f"Extracted {len(result['text'])} chars from MOL/SDF")
            
        except Exception as e:
            self._log(f"Error extracting MOL/SDF: {e}")
        
        return result

    # -------------------------
    # Main filter
    # -------------------------
    def inlet(self, body: dict, __user__: Optional[dict] = None) -> dict:
        self._log("=" * 60)
        self._log("INLET - Document Filter v1.1 (with ChemDraw/MOL/SDF)")
        
        if not self.valves.enabled:
            return body

        messages = body.get("messages", [])
        if not messages:
            return body

        files = self._extract_all_files(body, messages)
        if not files:
            self._log("No files found in request.")
            return body

        all_text_sections: List[str] = []
        all_images: List[Dict[str, Any]] = []

        for fobj in files:
            file_path = self._get_file_path(fobj)
            file_name = self._get_file_name(fobj)

            if not file_path or not os.path.exists(file_path):
                continue
            if not file_name:
                file_name = os.path.basename(file_path).lower()

            # Skip files handled by other filters
            if file_name.endswith((".ppt", ".pptx", ".pdf")):
                self._log(f"Skipping {file_name} (handled by PPT/PDF filter)")
                continue
            
            # Skip standalone images (handled by proxy)
            if file_name.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp")):
                self._log(f"Skipping {file_name} (standalone image, handled by proxy)")
                continue

            extracted = None
            
            # DOCX
            if file_name.endswith(".docx"):
                self._log(f"Processing DOCX: {file_name}")
                extracted = self._extract_docx(file_path)
            
            # XLSX
            elif file_name.endswith((".xlsx", ".xlsm")):
                self._log(f"Processing XLSX: {file_name}")
                extracted = self._extract_xlsx(file_path)
            
            # XLS (legacy)
            elif file_name.endswith(".xls"):
                self._log(f"Processing XLS: {file_name}")
                extracted = self._extract_xls(file_path)
            
            # CSV
            elif file_name.endswith(".csv"):
                self._log(f"Processing CSV: {file_name}")
                extracted = self._extract_csv(file_path, delimiter=',')
            
            # TSV
            elif file_name.endswith(".tsv"):
                self._log(f"Processing TSV: {file_name}")
                extracted = self._extract_csv(file_path, delimiter='\t')
            
            # TXT/MD/JSON
            elif file_name.endswith((".txt", ".md", ".json")):
                self._log(f"Processing text file: {file_name}")
                extracted = self._extract_text_file(file_path)
            
            # ChemDraw files
            elif file_name.endswith((".cdx", ".cdxml")):
                self._log(f"Processing ChemDraw: {file_name}")
                extracted = self._extract_chemdraw(file_path, file_name)
            
            # MOL/SDF structure files
            elif file_name.endswith((".mol", ".mol2", ".sdf")):
                self._log(f"Processing structure file: {file_name}")
                extracted = self._extract_mol_sdf(file_path, file_name)
            
            else:
                self._log(f"Unsupported file type: {file_name}")
                continue

            if extracted:
                if extracted.get("text"):
                    all_text_sections.append(f"=== {file_name} ===\n{extracted['text']}")
                if extracted.get("images"):
                    all_images.extend(extracted["images"])
                    self._log(f"Added {len(extracted['images'])} images from {file_name}")

        if not all_text_sections and not all_images:
            self._log("No extractable content produced.")
            return body

        # Build combined content
        last_message = messages[-1]
        original_prompt = self._extract_text_content(last_message.get("content", ""))

        combined_text = original_prompt.strip() + "\n\n"

        if all_text_sections:
            combined_text += "Document content (extracted):\n\n"
            combined_text += "\n\n".join(all_text_sections).strip()
            combined_text += "\n\n"

        if all_images:
            combined_text += f"Attached are {len(all_images)} images extracted from the documents.\n\n"

        # Build content blocks
        content_blocks: List[Dict[str, Any]] = [{"type": "text", "text": combined_text}]

        for img in all_images:
            if isinstance(img, dict) and img.get("type") == "image_url":
                iu = img.get("image_url", {})
                if isinstance(iu, dict) and iu.get("url"):
                    content_blocks.append({"type": "image_url", "image_url": {"url": iu["url"]}})

        messages[-1]["content"] = content_blocks
        body["messages"] = messages

        self._log(f"✅ Final payload:")
        self._log(f"   - Text sections: {len(all_text_sections)}")
        self._log(f"   - Image blocks: {len(all_images)}")
        self._log("=" * 60)
        return body

    def stream(self, event: dict, __user__: Optional[dict] = None) -> dict:
        return event

    def outlet(self, body: dict, __user__: Optional[dict] = None) -> dict:
        return body
